from fastapi import FastAPI, APIRouter, File, UploadFile, HTTPException, Form, Depends, Request, Body
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
try:
    from gtts import gTTS
except Exception:
    gTTS = None
import io
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import pymongo
import os
import logging


import os
import json
import logging
import pymongo
from datetime import datetime, timedelta
from bson import ObjectId
from fastapi import APIRouter, HTTPException

from fastapi import APIRouter, HTTPException
import pymongo
import os
import json
import logging
from pathlib import Path
from datetime import datetime, timezone
from typing import List, Dict, Any


# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional, Any
from dataclasses import dataclass
import uuid
from datetime import datetime, timezone, timedelta
import shutil
import httpx
import asyncio
from PyPDF2 import PdfReader
try:
    # pdfminer for more robust PDF text extraction
    from pdfminer.high_level import extract_text as pdfminer_extract_text
except Exception:
    pdfminer_extract_text = None
try:
    # python-docx for .docx extraction
    import docx
except Exception:
    docx = None
import hashlib
from passlib.context import CryptContext
import bcrypt
from jwt import encode as jwt_encode, decode as jwt_decode  # Import from PyJWT
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
import re
import secrets
import requests
import unicodedata
import tempfile
import json
import difflib
from io import BytesIO

import os
import re
import uuid
import logging
import pymongo
from fastapi import APIRouter, HTTPException
from PyPDF2 import PdfReader
from datetime import datetime, timezone





def strip_markdown(text: str) -> str:
    """Remove common Markdown formatting to produce clean plain text."""
    try:
        if not text:
            return text
        # Remove code fences ``` ```
        text = re.sub(r'```[\s\S]*?```', '', text)
        # Replace inline code `code` with code
        text = re.sub(r'`([^`]*)`', r'\1', text)
        # Remove ATX headings (# Headline)
        text = re.sub(r'^\s{0,3}#{1,6}\s*', '', text, flags=re.MULTILINE)
        # Replace bold/italic markers while preserving inner text
        text = re.sub(r'\*\*(.*?)\*\*', r"\1", text)
        text = re.sub(r'__(.*?)__', r"\1", text)
        text = re.sub(r'\*(.*?)\*', r"\1", text)
        text = re.sub(r'_(.*?)_', r"\1", text)
        # Remove list markers at start of lines
        text = re.sub(r'^[\s]*[-\*\+]\s+', '', text, flags=re.MULTILINE)
        # Remove remaining stray asterisks
        text = text.replace('*', '')
        # Collapse multiple blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
    except Exception:
        return text

# Authentication imports
from passlib.context import CryptContext
from datetime import timedelta
from fastapi.security import HTTPBearer

# Google OAuth configuration
GOOGLE_CLIENT_ID = os.environ.get('GOOGLE_CLIENT_ID')
GOOGLE_CLIENT_SECRET = os.environ.get('GOOGLE_CLIENT_SECRET')
GOOGLE_REDIRECT_URI = os.environ.get('GOOGLE_REDIRECT_URI', 'http://localhost:3000/auth/google/callback')

# Custom classes for message handling
@dataclass
class UserMessage:
    text: str
    file_contents: Optional[List[Any]] = None

@dataclass
class FileContentWithMimeType:
    file_path: str
    mime_type: str

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

logging.basicConfig(level=logging.INFO)
logging.info(f"GEMINI_API_KEY loaded: {'set' if os.environ.get('GEMINI_API_KEY') else 'NOT set'}")
if not os.environ.get('GEMINI_API_KEY'):
    logging.error("GEMINI_API_KEY is not set. Document analysis will not work!")

# Create the main app without a prefix
app = FastAPI()

# Serve static files from frontend build
static_path = ROOT_DIR.parent / "frontend" / "build" / "static"
if static_path.exists():
    app.mount("/static", StaticFiles(directory=str(static_path)), name="static")

# Serve frontend index.html for all non-API routes
@app.get("/{full_path:path}")
async def serve_frontend(full_path: str):
    if full_path.startswith("api/"):
        raise HTTPException(status_code=404, detail="API route not found")
        
    index_html = ROOT_DIR.parent / "frontend" / "build" / "index.html"
    if index_html.exists():
        return FileResponse(str(index_html))
    else:
        return {"error": "Frontend not built"}

# MongoDB helper function
def get_mongo_client():
    """Get a MongoDB client with connection verification"""
    mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
    try:
        client = pymongo.MongoClient(mongo_url, serverSelectionTimeoutMS=5000)
        # Force a connection to verify it works
        client.server_info()
        logging.info("MongoDB connection successful")
        return client
    except pymongo.errors.ServerSelectionTimeoutError as e:
        logging.error(f"MongoDB connection timeout: {str(e)}")
        raise HTTPException(status_code=500, detail="Database connection failed - timeout")
    except pymongo.errors.ConnectionFailure as e:
        logging.error(f"MongoDB connection failed: {str(e)}")
        raise HTTPException(status_code=500, detail="Database connection failed")
    except Exception as e:
        logging.error(f"MongoDB error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")

# MongoDB connection
try:
    mongo_url = os.environ.get('MONGO_URL')
    if not mongo_url:
        logging.warning("MONGO_URL not set - defaulting to localhost")
        mongo_url = 'mongodb://localhost:27017'
    
    # Configure MongoDB client with retry writes and proper timeouts
    client = AsyncIOMotorClient(
        mongo_url, 
        serverSelectionTimeoutMS=5000,
        heartbeatFrequencyMS=30000,  # Check every 30 seconds instead of 10
        retryWrites=True,
        connectTimeoutMS=30000,
        socketTimeoutMS=30000,
        maxIdleTimeMS=60000,
        maxPoolSize=10
    )
    db = client[os.environ.get('DB_NAME', 'legal_docs')]
    
    # Log connection details (without sensitive info)
    db_name = os.environ.get('DB_NAME', 'legal_docs')
    logging.info(f"MongoDB client initialized for database: {db_name}")
    logging.info("MongoDB connection will be tested on first use")
    
    # Add event handlers for connection states
    client.add_server_changed_listener(lambda event: 
        logging.info(f"MongoDB server state changed: {event.new_description.server_type}")
    )
except Exception as e:
    logging.error(f"Failed to initialize MongoDB client: {str(e)}")
    logging.error("Application will continue but database operations will fail")
    client = None
    db = None

# Add CORS middleware first. Make origins configurable via environment variables for easier dev/test workflows.
# - Set CORS_ALLOW_ALL=true to allow all origins (development only). This will set allow_credentials=False
#   because browsers disallow Access-Control-Allow-Credentials with wildcard origins.
# - Otherwise, set CORS_ORIGINS to a comma-separated list of allowed origins (e.g. "http://localhost:3000,http://127.0.0.1:3000").
cors_allow_all = os.environ.get('CORS_ALLOW_ALL', 'false').lower() in ('1', 'true', 'yes')
cors_origins_env = os.environ.get('CORS_ORIGINS')
if cors_allow_all:
    logging.warning('CORS_ALLOW_ALL is enabled: allowing all origins (credentials disabled). Do NOT enable in production.')
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_credentials=False,
        allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
        allow_headers=["Content-Type", "Authorization", "Accept", "Origin", "X-Requested-With"],
        expose_headers=["Content-Length", "Content-Range"],
        max_age=3600,
    )
else:
    if cors_origins_env:
        allow_origins = [o.strip() for o in cors_origins_env.split(',') if o.strip()]
    else:
        # sensible defaults for local development
        allow_origins = ["http://localhost:3000", "http://localhost:3001", "http://127.0.0.1:3000", "http://127.0.0.1:3001"]

    logging.info(f"CORS configured with allow_credentials=True and allow_origins={allow_origins}")
    app.add_middleware(
        CORSMiddleware,
        allow_credentials=True,
        allow_origins=allow_origins,
        allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
        allow_headers=["Content-Type", "Authorization", "Accept", "Origin", "X-Requested-With"],
        expose_headers=["Content-Length", "Content-Range"],
        max_age=3600,
    )

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Create uploads directory
UPLOAD_DIR = ROOT_DIR / "../uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

# Authentication setup
pwd_context = CryptContext(
    schemes=["bcrypt"],
    deprecated="auto",
    bcrypt__default_rounds=12,
    bcrypt__truncate_error=False  # This will allow truncation instead of raising an error
)
SECRET_KEY = os.environ.get('SECRET_KEY', secrets.token_urlsafe(32))
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 30

def hash_password(password: str) -> str:
    """Hash a password for storing using bcrypt directly.

    bcrypt enforces a 72-byte limit on the input. We truncate the UTF-8
    encoded bytes to 72 bytes, then hash using bcrypt.gensalt(). The resulting
    hash is ASCII-safe, so we store it as a utf-8 decoded string.
    """
    pw_bytes = password.encode('utf-8')
    if len(pw_bytes) > 72:
        logging.warning("Password exceeds 72 bytes, it will be truncated to 72 bytes before hashing")
        pw_bytes = pw_bytes[:72]

    hashed = bcrypt.hashpw(pw_bytes, bcrypt.gensalt(rounds=12))
    return hashed.decode('utf-8')

def create_access_token(data: dict, expires_delta: timedelta = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt_encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

def verify_token(token: str):
    """Verify a JWT access token and return the decoded payload as a dict.

    On success: returns the payload (dict).
    On failure: raises HTTPException(status_code=401) so calling routes receive an auth error.
    """
    if not token:
        raise HTTPException(status_code=401, detail="Authorization token missing")

    try:
        payload = jwt_decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        if not isinstance(payload, dict):
            raise HTTPException(status_code=401, detail="Invalid token payload")
        return payload
    except Exception as e:
        if "expired" in str(e).lower():
            logging.debug("JWT expired")
            raise HTTPException(status_code=401, detail="Token expired")
        logging.debug(f"JWT decode error: {e}")
        raise HTTPException(status_code=401, detail="Invalid authentication token")

# Pydantic Models
class Document(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    file_path: str
    file_type: str
    upload_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    analysis_status: str = "pending"  # pending, processing, completed, failed
    content: Optional[str] = None  # Extracted document text
    metadata: Optional[dict] = None  # Storage for content stats and extraction info
    summary: Optional[str] = None  # Executive summary
    key_points: Optional[List[dict]] = None  # Key points extracted from the document
    key_clauses: Optional[List[dict]] = None  # Important clauses found
    risk_assessment: Optional[dict] = None  # Structured risk assessment
    plain_english: Optional[str] = None  # Plain English explanation
    recommendations: Optional[List[dict]] = None  # Actionable recommendations
    analysis_version: str = "2.0"  # Track analysis system version

class DocumentCreate(BaseModel):
    filename: str
    file_type: str

class QuestionRequest(BaseModel):
    document_id: str
    question: str
    session_id: Optional[str] = None

class ChatMessage(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str
    session_id: str
    question: str
    answer: str
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class GlobalChatMessage(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    user_id: Optional[str] = None  # None for anonymous users
    chat_mode: str = "general"  # "general" or "document"
    selected_document: Optional[str] = None
    question: str
    answer: str
    document_metadata: Optional[dict] = None  # Store relevant document context
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class AnalysisRequest(BaseModel):
    document_id: str

class ExportRequest(BaseModel):
    sections: List[str]

# User models for authentication
class User(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: str
    name: str
    google_id: Optional[str] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class LoginRequest(BaseModel):
    email: str
    password: str

class RegisterRequest(BaseModel):
    email: str
    password: str
    name: str

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: dict

# Document text extraction system
def validate_pdf(file_path: str) -> bool:
    """Validate PDF file integrity"""
    try:
        with open(file_path, 'rb') as file:
            # Check PDF header
            header = file.read(5)
            if header != b'%PDF-':
                return False
            
            # Try to load with PyPDF2 to check structure
            reader = PdfReader(file_path)
            if len(reader.pages) == 0:
                return False
                
            return True
    except Exception as e:
        logging.warning(f"PDF validation failed: {e}")
        return False

def extract_document_text(file_path: str, file_type: str) -> tuple[str, str]:
    """
    Enhanced document text extraction with multiple fallback methods.
    Returns: (extracted_text, method_used)
    """
    logging.info(f"Starting document extraction for {file_path} of type {file_type}")
    
    text = ""
    method = ""
    error_messages = []
    
    # For PDFs, validate first
    if file_type == 'application/pdf' or file_path.lower().endswith('.pdf'):
        if not validate_pdf(file_path):
            error_messages.append("Invalid or corrupted PDF file")
            return "Invalid or corrupted PDF file", "failed"
    
    # PDF Handling
    if file_type == 'application/pdf' or file_path.lower().endswith('.pdf'):
        # 1. Try pdfminer.six first (more robust)
        try:
            from pdfminer.high_level import extract_text
            from pdfminer.layout import LAParams
            
            # Configure LAParams for better text extraction
            laparams = LAParams(
                line_margin=0.5,
                word_margin=0.1,
                char_margin=2.0,
                boxes_flow=0.5,
                detect_vertical=True
            )
            
            pdfminer_text = extract_text(file_path, laparams=laparams)
            if pdfminer_text and pdfminer_text.strip():
                text = pdfminer_text
                # Clean up extracted text
                text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
                text = re.sub(r'\n\s*\n', '\n\n', text)  # Fix paragraph breaks
                text = text.strip()
                
                if len(text) > 100:  # Check if we got meaningful text
                    method = "pdfminer"
                    logging.info(f"pdfminer extraction successful: {len(text)} chars")
                    return text, method
        except Exception as e:
            error_messages.append(f"pdfminer failed: {str(e)}")
            
        # 2. Try PyPDF2 as backup
        try:
            reader = PdfReader(file_path)
            pages = []
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        pages.append(page_text)
                        logging.info(f"PyPDF2 extracted page {i}: {len(page_text)} chars")
                except Exception as e:
                    logging.warning(f"PyPDF2 failed on page {i}: {e}")
            
            pypdf_text = "\n".join(pages)
            if pypdf_text.strip() and len(pypdf_text) > 100:
                text = pypdf_text
                # Clean up extracted text
                text = re.sub(r'\s+', ' ', text)
                text = re.sub(r'\n\s*\n', '\n\n', text)
                text = text.strip()
                method = "PyPDF2"
                logging.info(f"PyPDF2 extraction successful: {len(text)} chars")
                return text, method
        except Exception as e:
            error_messages.append(f"PyPDF2 failed: {str(e)}")
            
        # 2. Try pdfminer.six
        try:
            if pdfminer_extract_text:
                pdfminer_text = pdfminer_extract_text(file_path)
                if pdfminer_text and len(pdfminer_text.strip()) > len(text.strip()):
                    text = pdfminer_text
                    method = "pdfminer"
                    logging.info(f"pdfminer extraction successful: {len(text)} chars")
                    return text, method
        except Exception as e:
            error_messages.append(f"pdfminer failed: {str(e)}")
            
        # 3. Try OCR as last resort
        if len(text.strip()) < 100:
            try:
                import pytesseract
                from PIL import Image
                from pdf2image import convert_from_path
                import os
                
                logging.info("Attempting OCR extraction")
                ocr_text = []
                
                # Configure tesseract for better accuracy
                custom_config = r'--oem 3 --psm 6'
                
                with tempfile.TemporaryDirectory() as tempdir:
                    try:
                        # Convert PDF to images with higher DPI for better OCR
                        images = convert_from_path(file_path, output_folder=tempdir, dpi=300)
                        
                        for i, image in enumerate(images):
                            # Preprocess image for better OCR
                            # Convert to RGB if not already
                            if image.mode != 'RGB':
                                image = image.convert('RGB')
                            
                            # Try OCR with custom config
                            page_text = pytesseract.image_to_string(image, config=custom_config)
                            
                            if page_text.strip():
                                # Clean up OCR text
                                cleaned_text = re.sub(r'\s+', ' ', page_text)  # Normalize whitespace
                                cleaned_text = re.sub(r'\n\s*\n', '\n\n', cleaned_text)  # Fix paragraph breaks
                                ocr_text.append(cleaned_text.strip())
                                logging.info(f"OCR extracted page {i}: {len(cleaned_text)} chars")
                    except Exception as e:
                        logging.error(f"Error during PDF to image conversion: {e}")
                        
                final_ocr_text = "\n\n".join(ocr_text)
                if final_ocr_text and len(final_ocr_text.strip()) > len(text.strip()):
                    text = final_ocr_text
                    method = "OCR"
                    logging.info(f"OCR extraction successful: {len(text)} chars")
                    return text, method
            except Exception as e:
                error_messages.append(f"OCR failed: {str(e)}")
                
    # DOCX Handling
    elif file_path.lower().endswith('.docx'):
        try:
            if docx:
                doc = docx.Document(file_path)
                paragraphs = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        paragraphs.append(p.text)
                text = "\n".join(paragraphs)
                if text.strip():
                    method = "python-docx"
                    logging.info(f"DOCX extraction successful: {len(text)} chars")
                    return text, method
        except Exception as e:
            error_messages.append(f"DOCX extraction failed: {str(e)}")
    
    # Text file handling
    else:
        # Try UTF-8 first
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                if text.strip():
                    method = "utf-8"
                    logging.info(f"UTF-8 read successful: {len(text)} chars")
                    return text, method
        except UnicodeDecodeError:
            # Try latin-1 as fallback
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                    if text.strip():
                        method = "latin-1"
                        logging.info(f"Latin-1 read successful: {len(text)} chars")
                        return text, method
            except Exception as e:
                error_messages.append(f"Text read failed: {str(e)}")
    
    # If we reach here, all methods failed or text is empty
    if not text.strip():
        error_detail = "; ".join(error_messages) if error_messages else "No valid text content found"
        logging.error(f"All extraction methods failed: {error_detail}")
        return f"Document text extraction failed: {error_detail}", "failed"
    
    return text, method

def extract_text_pdf_fallback(file_path: str) -> str:
    """Attempt to extract PDF text using pdfminer.six as a fallback."""
    try:
        if pdfminer_extract_text:
            logging.info(f"Attempting pdfminer.six extraction on {file_path}")
            try:
                text = pdfminer_extract_text(file_path)
                logging.info(f"pdfminer.six extraction result length: {len(text) if text else 0}")
                if text:
                    logging.debug(f"pdfminer.six sample: {text[:200]}")
                return text or ""
            except Exception as e:
                logging.warning(f"pdfminer extraction failed: {e}")
        else:
            logging.info("pdfminer.six not available in environment")
    except Exception as e:
        logging.error(f"Unexpected pdfminer fallback error: {e}")
    return ""

def extract_text_docx(file_path: str) -> str:
    """Extract text from .docx files using python-docx."""
    try:
        if docx:
            logging.info(f"Attempting python-docx extraction for .docx: {file_path}")
            try:
                document = docx.Document(file_path)
                paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
                logging.info(f"python-docx extraction result length: {sum(len(p) for p in paragraphs)}")
                if paragraphs:
                    logging.debug(f"python-docx sample: {paragraphs[0][:200]}")
                return "\n".join(paragraphs)
            except Exception as e:
                logging.warning(f"python-docx extraction failed: {e}")
        else:
            logging.info("python-docx not available in environment")
    except Exception as e:
        logging.error(f"Unexpected docx fallback error: {e}")
    return ""

GEMINI_API_KEY = os.getenv('GEMINI_API_KEY', 'AIzaSyCuBibZH55rvDR8b8Utty9ThK_hUTBh3Es')
GEMINI_API_URL = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={GEMINI_API_KEY}"

async def send_message(prompt: str, mode: str = "general", context: str = "") -> str:
    logging.debug(f"Starting AI request with prompt length: {len(prompt)}, mode: {mode}")
    
    # Set up the system message based on mode
    if mode == "document":
        system_prefix = """You are a legal document analysis assistant. Your task is to:
1. Answer questions based ONLY on the provided document content
2. If information isn't in the document, clearly state that
3. Quote relevant sections when possible
4. Keep responses concise and focused
5. Do not make assumptions beyond the document content

"""
        full_prompt = system_prefix + prompt
    else:
        full_prompt = prompt

    async def local_fallback(p: str) -> str:
        """Generate a safe, concise fallback response when the external AI is unavailable.

        This provides reasonable behavior for development and offline testing so the
        `/api/chat` and document analysis endpoints remain functional.
        """
        # If the prompt looks like a long document, return a short structured analysis
        text = p.strip()
        if len(text) > 2000 or '\n' in text and len(text) > 400:
            # Very short heuristic analysis: executive summary + key points
            summary = text[:800].strip()
            return (
                "EXECUTIVE SUMMARY: This document appears to be a legal text. "
                "A quick read indicates the main topics and obligations are: "
                f"{summary[:300]}...\n\nKEY POINTS: (1) Check parties and dates; (2) Look for termination and liability clauses; (3) Identify obligations and deliverables."
            )

        # If it's a question-style prompt, reply concisely
        if text.endswith('?') or text.lower().startswith(('what', 'how', 'why', 'explain', 'when', 'where', 'who')):
            return "I'm a local fallback assistant: I can give quick answers, but the full AI service isn't available. Try a concise question or enable the AI API for detailed responses."

        # Default echo-style friendly reply
        short = text[:400]
        return f"Acknowledged. Here's a brief response based on your input: {short}{'...' if len(text) > 400 else ''}"

    # Try external Gemini API if configured; otherwise use local fallback
    if GEMINI_API_KEY:
        max_retries = 2
        for attempt in range(max_retries):
            try:
                async with httpx.AsyncClient(timeout=30.0) as client:
                    response = await client.post(
                        GEMINI_API_URL,
                        json={
                            "contents": [{"parts": [{"text": prompt}]}]
                        },
                        headers={"Content-Type": "application/json"},
                    )
                    if response.status_code == 200:
                        data = response.json()
                        logging.debug(f"Gemini API response keys: {list(data.keys()) if isinstance(data, dict) else type(data)}")
                        try:
                            candidate = data.get('candidates', [])[0]
                            response_text = candidate.get('content', {}).get('parts', [])[0].get('text', '')
                            if response_text:
                                return response_text
                        except Exception:
                            logging.debug('Unexpected Gemini response structure, falling back')
                    else:
                        logging.warning(f"Gemini status {response.status_code}; body: {response.text}")
            except Exception as e:
                logging.warning(f"Gemini API attempt {attempt+1} failed: {e}")

    # External API disabled or failed -> return local fallback
    logging.info('Using local AI fallback response')
    # local_fallback is async; await it directly instead of passing the coroutine into
    # run_in_executor (which expects a synchronous callable). Awaiting guarantees a
    # concrete string is returned and avoids "coroutine was never awaited" warnings.
    return await local_fallback(prompt)

@api_router.get("/")
async def root():
    return {"message": "Legal Document AI Assistant API"}

@api_router.post("/auth/login")
async def login(request: LoginRequest):
    """Login or auto-register user and send verification if needed"""
    try:
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        import logging

        logging.info(f"Login attempt for email: {request.email}")

        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        user = sync_db.users.find_one({"email": request.email})
        logging.info(f"Found user: {user is not None}")
        if user:
            logging.info(f"User ID: {user.get('id')}")
            logging.info(f"Email verified: {user.get('email_verified')}")
            
        stored_password = user.get('password', '') if user else ''

        # 🧩 If user does not exist → auto-register and send verification code
        if not user:
            code = secrets.token_hex(3)
            expires = (datetime.utcnow() + timedelta(minutes=15)).isoformat()

            hashed_pw = bcrypt.hashpw(request.password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

            new_user = {
                "email": request.email,
                "password": hashed_pw,
                "email_verified": False,
                "email_verify_code": code,
                "email_verify_expires": expires,
                "created_at": datetime.utcnow().isoformat(),
            }

            inserted = sync_db.users.insert_one(new_user)
            user_id = str(inserted.inserted_id)

            # --- send verification email ---
            smtp_host = os.environ.get('SMTP_HOST')
            smtp_port = int(os.environ.get('SMTP_PORT', '587'))
            smtp_user = os.environ.get('SMTP_USER')
            smtp_pass = os.environ.get('SMTP_PASS')
            from_email = os.environ.get('SMTP_FROM', 'CovenantAI <noreply@legaldocai.com>')

            if smtp_host and smtp_user and smtp_pass:
                try:
                    message = MIMEMultipart("alternative")
                    message['From'] = from_email
                    message['To'] = request.email
                    message['Subject'] = "Verify Your Email - CovenantAI"

                    text_body = f"""Hi there,

Please verify your email using this code: {code}

This code will expire in 15 minutes.

Best,
CovenantAI Team
"""

                    html_body = f"""
                    <div style="font-family: 'Segoe UI', sans-serif; background-color: #f4f4f7; padding: 20px;">
                        <div style="max-width: 600px; margin: auto; background: #fff; border-radius: 10px;
                                    box-shadow: 0 3px 8px rgba(0,0,0,0.05); overflow: hidden;">
                            <div style="background: linear-gradient(135deg, #3b82f6, #2563eb); color: white; padding: 18px; text-align: center;">
                                <h2 style="margin: 0;">Verify Your Email</h2>
                            </div>
                            <div style="padding: 25px; color: #333;">
                                <p>Hi <strong>{request.email}</strong>,</p>
                                <p>Thanks for joining <b>CovenantAI</b>! Please verify your email using the code below:</p>
                                <div style="text-align: center; margin: 30px 0;">
                                    <div style="display: inline-block; background: #2563eb; color: white; 
                                                font-size: 20px; letter-spacing: 3px; padding: 12px 24px; 
                                                border-radius: 6px;">
                                        {code}
                                    </div>
                                </div>
                                <p>This code will expire in <b>15 minutes</b>.</p>
                                <p>If you didn’t create an account, ignore this email.</p>
                                <hr style="border:none; border-top:1px solid #eee; margin: 25px 0;">
                                <p style="font-size: 13px; color: #777;">© {datetime.utcnow().year} CovenantAI. All rights reserved.</p>
                            </div>
                        </div>
                    </div>
                    """

                    message.attach(MIMEText(text_body, "plain"))
                    message.attach(MIMEText(html_body, "html"))

                    with smtplib.SMTP(smtp_host, smtp_port) as server:
                        server.starttls()
                        server.login(smtp_user, smtp_pass)
                        server.send_message(message)

                    logging.info(f"Verification email sent to {request.email}")
                except Exception as mail_error:
                    logging.error(f"Failed to send verification email: {mail_error}")
            else:
                logging.info(f"SMTP not configured, verification code for {request.email}: {code}")

            sync_client.close()
            return {
                "message": "Verification code sent. Please verify your email before login.",
                "email": request.email,
                "requires_verification": True
            }

        # 🧩 Existing user → verify password
        try:
            logging.info(f"Verifying password for user: {request.email}")
            logging.info(f"Stored password: {stored_password[:20] if stored_password else 'EMPTY'}...")
            logging.info(f"Input password: {request.password[:5]}...")
            
            # Add debug logging
            logging.info("Password verification details:")
            logging.info(f"Input password length: {len(request.password)}")
            logging.info(f"Stored hash length: {len(stored_password)}")
            logging.info(f"Stored hash type: {type(stored_password)}")
            logging.info(f"Full user record: {user}")
            logging.info(f"Checking bcrypt password...")

            # Prepare bytes for bcrypt
            password_bytes = request.password.encode('utf-8')
            hash_bytes = stored_password.encode('utf-8') if isinstance(stored_password, str) else stored_password
            
            logging.info(f"Password bytes: {password_bytes[:20]}...")
            logging.info(f"Hash bytes: {hash_bytes[:20]}...")

            valid = bcrypt.checkpw(password_bytes, hash_bytes)
            logging.info(f"Password valid: {valid}")

            if not valid:
                logging.error(f"Password check failed for {request.email}")
                sync_client.close()
                raise HTTPException(status_code=401, detail="Invalid credentials")

        except HTTPException:
            # Re-raise HTTP exceptions
            raise
        except Exception as e:
            logging.error(f"Password verification failed: {str(e)}", exc_info=True)
            sync_client.close()
            raise HTTPException(status_code=401, detail="Invalid credentials")

        # 🧩 If user email not verified → resend code if needed
        if not user.get('email_verified', False):
            code = user.get('email_verify_code') or secrets.token_hex(3)
            expires = (datetime.utcnow() + timedelta(minutes=15)).isoformat()

            sync_db.users.update_one(
                {"_id": user["_id"]},
                {"$set": {
                    "email_verify_code": code,
                    "email_verify_expires": expires
                }}
            )

            # Re-send verification email
            smtp_host = os.environ.get('SMTP_HOST')
            smtp_port = int(os.environ.get('SMTP_PORT', '587'))
            smtp_user = os.environ.get('SMTP_USER')
            smtp_pass = os.environ.get('SMTP_PASS')
            from_email = os.environ.get('SMTP_FROM', 'CovenantAI <noreply@legaldocai.com>')

            if smtp_host and smtp_user and smtp_pass:
                try:
                    message = MIMEMultipart("alternative")
                    message['From'] = from_email
                    message['To'] = user['email']
                    message['Subject'] = "Verify Your Email - CovenantAI"

                    text_body = f"""Hi {user['email']},

Please verify your email using this code: {code}

This code will expire in 15 minutes.

Best,
CovenantAI Team
"""

                    html_body = f"""
                    <div style="font-family: 'Segoe UI', sans-serif; background-color: #f4f4f7; padding: 20px;">
                        <div style="max-width: 600px; margin: auto; background: #fff; border-radius: 10px;
                                    box-shadow: 0 3px 8px rgba(0,0,0,0.05); overflow: hidden;">
                            <div style="background: linear-gradient(135deg, #3b82f6, #2563eb); color: white; padding: 18px; text-align: center;">
                                <h2 style="margin: 0;">Verify Your Email</h2>
                            </div>
                            <div style="padding: 25px; color: #333;">
                                <p>Hi <strong>{user['email']}</strong>,</p>
                                <p>Use the code below to verify your account:</p>
                                <div style="text-align: center; margin: 30px 0;">
                                    <div style="display: inline-block; background: #2563eb; color: white; 
                                                font-size: 20px; letter-spacing: 3px; padding: 12px 24px; 
                                                border-radius: 6px;">
                                        {code}
                                    </div>
                                </div>
                                <p>This code will expire in <b>15 minutes</b>.</p>
                                <p>If you didn’t request this, ignore this email.</p>
                                <hr style="border:none; border-top:1px solid #eee; margin: 25px 0;">
                                <p style="font-size: 13px; color: #777;">© {datetime.utcnow().year} CovenantAI. All rights reserved.</p>
                            </div>
                        </div>
                    </div>
                    """

                    message.attach(MIMEText(text_body, "plain"))
                    message.attach(MIMEText(html_body, "html"))

                    with smtplib.SMTP(smtp_host, smtp_port) as server:
                        server.starttls()
                        server.login(smtp_user, smtp_pass)
                        server.send_message(message)

                    logging.info(f"Verification email re-sent to {user['email']}")
                except Exception as mail_error:
                    logging.error(f"Failed to re-send verification email: {mail_error}")
            else:
                logging.info(f"SMTP not configured, code for {user['email']}: {code}")

            sync_client.close()
            return {
                "message": "Email verification required",
                "email": user['email'],
                "requires_verification": True
            }

        # 🧩 If verified → create access token and login
        user_id = user.get('id') or str(user.get('_id'))  # Fallback to string _id if no id field
        logging.info(f"Creating access token for user ID: {user_id}")
        
        token_data = {
            "sub": user['email'], 
            "user_id": user_id,
            "email": user['email']
        }
        access_token = create_access_token(
            data=token_data,
            expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        )
        logging.info(f"Access token created successfully")

        sync_client.close()

        # Create standardized user response
        user_response = {
            "id": user_id,
            "email": user['email'],
            "name": user.get('name', '')
        }
        logging.info(f"Returning successful login response")
        
        return TokenResponse(
            access_token=access_token,
            user=user_response
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Login error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Login failed")

@api_router.get("/history")
async def get_history(request: Request):
    """Get user's document history"""
    try:
        # Extract token from Authorization header
        auth_header = request.headers.get("Authorization")
        if not auth_header or not auth_header.startswith("Bearer "):
            raise HTTPException(status_code=401, detail="Authorization header missing or invalid")

        token = auth_header.replace("Bearer ", "")
        payload = verify_token(token)
        if not payload:
            raise HTTPException(status_code=401, detail="Invalid token")

        user_id = payload.get("user_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token payload")

        # Get user's documents
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        documents = list(sync_db.documents.find({"user_id": user_id}).sort("created_at", -1))
        sync_client.close()

        # Format documents for response
        history = []
        for doc in documents:
            history.append({
                'id': doc['id'],
                'filename': doc.get('filename', 'Untitled Document'),
                'document_text': doc.get('document_text', '')[:200] + '...' if len(doc.get('document_text', '')) > 200 else doc.get('document_text', ''),
                'analysis_result': doc.get('analysis_result', {}),
                'summary': doc.get('summary', ''),
                'risk_score': doc.get('risk_score', {}),
                'critical_flags': doc.get('critical_flags', []),
                'analysis_status': doc.get('analysis_status', 'completed'),
                'created_at': doc.get('created_at', '').isoformat() if isinstance(doc.get('created_at'), datetime) else str(doc.get('created_at', ''))
            })

        return history

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Get history error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve history")




def safe_json(obj):
    """Recursively make MongoDB objects (ObjectId, datetime, etc.) JSON serializable"""
    if isinstance(obj, ObjectId):
        return str(obj)
    elif isinstance(obj, datetime):
        return obj.isoformat()
    elif isinstance(obj, dict):
        return {k: safe_json(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [safe_json(v) for v in obj]
    return obj


@api_router.get("/history/guest")
async def get_guest_history():
    """Get guest document history (documents with no user_id)"""
    try:
        # --- 1. Connect to MongoDB ---
        mongo_url = os.environ.get("MONGO_URL", "mongodb://localhost:27017")
        db_name = os.environ.get("DB_NAME", "legal_docs")
        sync_client = pymongo.MongoClient(mongo_url)
        sync_db = sync_client[db_name]

        # --- 2. Fetch guest (unauthenticated) documents ---
        documents = list(sync_db.documents.find({
            "$or": [
                {"user_id": {"$exists": False}},
                {"user_id": None},
                {"user_id": ""}
            ]
        }).sort("upload_date", -1))  # Sort newest first

        # --- 3. Close client ---
        sync_client.close()

        # --- 4. Format for response ---
        history = []
        for doc in documents:
            # Handle missing analysis_result gracefully
            analysis_result = doc.get("analysis_result")
            if not analysis_result:
                raw_summary = doc.get("summary") or ""
                analysis_result = {}
                if isinstance(raw_summary, str) and raw_summary.strip().startswith("{"):
                    try:
                        analysis_result = json.loads(raw_summary)
                    except Exception:
                        analysis_result = {"raw_text": raw_summary}
                else:
                    analysis_result = {"raw_text": raw_summary}

            # Normalize upload_date
            upload_date_val = doc.get("upload_date")
            upload_dt = None
            if isinstance(upload_date_val, datetime):
                upload_dt = upload_date_val
            elif isinstance(upload_date_val, str):
                try:
                    upload_dt = datetime.fromisoformat(upload_date_val)
                except Exception:
                    upload_dt = None

            # Compute expiry time (10 minutes after upload)
            expires_at = (upload_dt + timedelta(minutes=10)).isoformat() if upload_dt else None
            upload_time_str = upload_dt.isoformat() if upload_dt else (
                str(upload_date_val) if upload_date_val is not None else ""
            )

            # Append cleaned record
            history.append({
                "id": str(doc.get("_id", "")),
                "filename": doc.get("filename", "Untitled Document"),
                "analysis_result": safe_json(analysis_result),
                "upload_time": upload_time_str,
                "expires_at": expires_at
            })

        return history

    except Exception as e:
        import traceback
        logging.error(f"Get guest history error: {e}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve guest history: {str(e)}")



@api_router.post("/history")
async def save_to_history(request: Request, data: dict = Body(...)):
    """Save document to history"""
    try:
        logging.info(f"🔵 /history POST endpoint called")
        logging.info(f"   Request data keys: {list(data.keys())}")
        
        # Extract token from Authorization header
        auth_header = request.headers.get("Authorization")
        logging.info(f"   Authorization header: {auth_header[:20] if auth_header else 'MISSING'}...")
        
        if not auth_header or not auth_header.startswith("Bearer "):
            logging.error("Authorization header missing or invalid")
            raise HTTPException(status_code=401, detail="Authorization header missing or invalid")

        token = auth_header.replace("Bearer ", "")
        payload = verify_token(token)
        if not payload:
            logging.error("Invalid token")
            raise HTTPException(status_code=401, detail="Invalid token")

        user_id = payload.get("user_id")
        if not user_id:
            logging.error("Invalid token payload - no user_id")
            raise HTTPException(status_code=401, detail="Invalid token payload")

        logging.info(f"   User ID: {user_id}")

        document_text = data.get('documentText')
        analysis_result = data.get('analysisResult', {})
        filename = data.get('filename', 'Untitled Document')

        logging.info(f"   Document text length: {len(document_text) if document_text else 0}")
        logging.info(f"   Analysis result keys: {list(analysis_result.keys()) if analysis_result else []}")
        logging.info(f"   Filename: {filename}")

        if not document_text:
            logging.error("Document text is required")
            raise HTTPException(status_code=400, detail="Document text is required")

        if not analysis_result:
            logging.error("Analysis result is required")
            raise HTTPException(status_code=400, detail="Analysis result is required")

        # Save to database
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        doc_data = {
            "id": str(uuid.uuid4()),
            "user_id": user_id,
            "filename": filename,
            "document_text": document_text,
            "analysis_result": analysis_result,
            "created_at": datetime.now(timezone.utc),
            "summary": analysis_result.get('summary', ''),
            "risk_score": analysis_result.get('riskScore', {}),
            "critical_flags": analysis_result.get('criticalFlags', []),
            "analysis_status": "completed"
        }

        logging.info(f"   Inserting document: {doc_data.get('id')}")
        sync_db.documents.insert_one(doc_data)
        sync_client.close()
        
        logging.info(f"✅ Document saved to history successfully: {doc_data.get('id')}")

        return {"message": "Document saved to history"}

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Save to history error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to save to history")

@api_router.post("/chat")
async def chat_endpoint(request: Request, data: dict = Body(...)):
    """Chat with AI about legal questions - works for authenticated and anonymous users"""
    logging.info(f"🤖 CHAT ENDPOINT CALLED - Data: {data}")
    sync_client = None
    try:
        # Extract token from Authorization header (optional for anonymous users)
        user_id = None
        auth_header = request.headers.get("Authorization")
        logging.info(f"🤖 Auth header: {auth_header if auth_header else 'NONE'}")
        if auth_header and auth_header.startswith("Bearer "):
            logging.info(f"🤖 Bearer token found, extracting...")
            token = auth_header.replace("Bearer ", "")
            logging.info(f"🤖 Token: {token[:50]}...")
            try:
                payload = verify_token(token)
                logging.info(f"🤖 verify_token returned: {payload}")
                if payload:
                    user_id = payload.get("user_id")
                    logging.info(f"🤖 User ID extracted: {user_id}")
                else:
                    logging.info(f"🤖 verify_token returned None/False")
            except Exception as e:
                logging.error(f"🤖 verify_token EXCEPTION: {e}", exc_info=True)

        message = data.get('message')
        context = data.get('context', [])
        mode = data.get('mode', 'general')  # Get chat mode from request
        selected_document = data.get('selected_document')

        if not message:
            raise HTTPException(status_code=400, detail="Message is required")

        # Create MongoDB client for document retrieval
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        # Build context from previous messages and document if in document mode
        document_content_for_ai = ""
        if mode == "document" and selected_document:
            try:
                # Get document content from database
                logging.info(f"Fetching document {selected_document} for chat")
                document = sync_db.documents.find_one({"id": selected_document})
                if not document:
                    raise HTTPException(status_code=404, detail="Document not found")

                # Try different content fields in order of preference
                document_content_for_ai = (
                    document.get('content', '') or 
                    document.get('document_text', '') or
                    document.get('summary', '')
                )

                # Fallback: If no content, try to read from file
                if not document_content_for_ai:
                    file_path = document.get('file_path')
                    file_type = document.get('file_type', '')
                    if file_path:
                        from pathlib import Path
                        # Try the file_path as stored
                        if Path(file_path).exists():
                            try:
                                if file_type == 'application/pdf' or str(file_path).lower().endswith('.pdf'):
                                    # Extract from PDF
                                    from PyPDF2 import PdfReader
                                    reader = PdfReader(file_path)
                                    pages = []
                                    for page in reader.pages:
                                        text = page.extract_text()
                                        if text:
                                            pages.append(text)
                                    document_content_for_ai = "\n".join(pages)
                                    logging.info(f"Extracted PDF content for chat, length={len(document_content_for_ai)}")
                                else:
                                    # Text file
                                    with open(file_path, 'r', encoding='utf-8') as f:
                                        document_content_for_ai = f.read()
                                    logging.info(f"Loaded text content from file for chat, length={len(document_content_for_ai)}")
                            except Exception as e:
                                logging.error(f"Failed to read file for chat fallback: {e}")

                if not document_content_for_ai:
                    logging.error(f"Document {selected_document} has no readable content")
                    raise HTTPException(status_code=400, detail="Document has no analyzable content")

                logging.info(f"Found document with content length: {len(document_content_for_ai)}")

            except HTTPException:
                raise
            except Exception as e:
                logging.error(f"Error accessing document {selected_document}: {e}")
                raise HTTPException(status_code=500, detail="Failed to access document")
        
        # Build system message
        if mode == "document" and selected_document and document_content_for_ai:
            system_message = (
                "You are a legal AI assistant analyzing a specific document. "
                "Answer questions based ONLY on the content of this document. "
                "Reference specific sections when possible. "
                "If information is not found in the document, clearly state that. "
                "Do not provide general legal advice beyond what is in the document.\n\n"
                "DOCUMENT CONTENT FOR REFERENCE:\n"
                "---START OF DOCUMENT---\n"
                f"{document_content_for_ai}\n"
                "---END OF DOCUMENT---\n\n"
            )
        else:
            system_message = (
                "You are a legal AI assistant. Provide helpful, accurate legal information "
                "and analysis. Note: This is general legal information and not formal legal "
                "advice. Always consult with a qualified attorney for your specific situation. "
                "Keep responses concise and direct."
            )

        messages = [{"role": "system", "content": system_message}]

        # Add recent context (last 5 messages)
        for msg in context[-5:]:
            messages.append({"role": msg['role'], "content": msg['content']})

        messages.append({"role": "user", "content": message})

        # Get AI response
        if mode == "document" and selected_document and document_content_for_ai:
            # Create a more focused prompt for document-specific questions
            doc_prompt = f"""Based on the document content provided in your system message, answer the following question.
Answer ONLY based on what is actually in the document. If the answer is not found in the document, clearly state that.

USER QUESTION: {message}

Remember: Base your entire answer solely on the document content. Quote relevant sections when applicable."""
            
            response = await send_message(doc_prompt, mode="document", context=document_content_for_ai)
        else:
            response = await send_message(message)
            
        if not response or len(response.strip()) < 10:
            raise Exception("Empty or invalid response from AI service")
        # Strip markdown and normalize spacing
        response = strip_markdown(response)

        # Limit to approximately 100 words and make it crisp
        words = response.split()
        if len(words) > 100:
            response = ' '.join(words[:100]) + '...'
        # Make it more concise by removing unnecessary phrases
        response = re.sub(r'\s+', ' ', response)  # Multiple spaces to single
        response = response.strip()

        # Save to global chat history for all users (authenticated and anonymous)
        session_id = data.get('session_id', f"global_{uuid.uuid4()}")
        
        # For document mode, store the relevant context
        document_metadata = None
        if mode == "document" and selected_document:
            document = sync_db.documents.find_one({"id": selected_document})
            if document:
                document_metadata = {
                    "filename": document.get("filename"),
                    "summary": document.get("summary", "")[:500],  # Store a snippet of the summary
                    "content_length": len(document.get("content", "") or document.get("document_text", ""))
                }
        
        global_chat_message = GlobalChatMessage(
            session_id=session_id,
            user_id=user_id,
            chat_mode=mode,
            selected_document=selected_document,
            question=message,
            answer=response,
            document_metadata=document_metadata
        )

        global_chat_dict = global_chat_message.model_dump()
        global_chat_dict['timestamp'] = global_chat_dict['timestamp'].isoformat()

        # Save to chat history using the existing client
        sync_db.global_chat_messages.insert_one(global_chat_dict)

        return {"response": response, "session_id": session_id}

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"❌ Chat error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Chat failed")
    finally:
        if sync_client:
            try:
                sync_client.close()
            except Exception:
                pass

@api_router.post("/tts")
async def text_to_speech(data: dict = Body(...)):
    """Converts text to speech and returns an audio file."""
    try:
        text = data.get("text")
        if not text:
            raise HTTPException(status_code=400, detail="Text is required")

        # gTTS is optional; if it's not installed, return a clear error so the server
        # can still run in environments without that dependency.
        if gTTS is None:
            raise HTTPException(status_code=501, detail="Text-to-speech not available: gTTS dependency not installed")

        # Create gTTS object
        tts = gTTS(text=text, lang='en', slow=False)

        # Save to a in-memory file
        mp3_fp = io.BytesIO()
        tts.write_to_fp(mp3_fp)
        mp3_fp.seek(0)

        return StreamingResponse(mp3_fp, media_type="audio/mpeg")

    except Exception as e:
        logging.error(f"TTS error: {e}")
        raise HTTPException(status_code=500, detail="Text-to-speech conversion failed")

@api_router.get("/analytics")
async def get_analytics_endpoint(request: Request):
    """Get user analytics"""
    try:
        # Extract token from Authorization header
        auth_header = request.headers.get("Authorization")
        if not auth_header or not auth_header.startswith("Bearer "):
            raise HTTPException(status_code=401, detail="Authorization header missing or invalid")

        token = auth_header.replace("Bearer ", "")
        payload = verify_token(token)
        if not payload:
            raise HTTPException(status_code=401, detail="Invalid token")

        user_id = payload.get("user_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token payload")

        # Get analytics data
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        # Get total documents
        total_docs = sync_db.documents.count_documents({"user_id": user_id})

        # Get documents from this month
        current_month = datetime.now().strftime('%Y-%m')
        this_month_count = sync_db.documents.count_documents({
            "user_id": user_id,
            "created_at": {"$gte": datetime.now().replace(day=1, hour=0, minute=0, second=0, microsecond=0)}
        })

        # Get average risk score and other analytics
        documents = list(sync_db.documents.find({"user_id": user_id}))
        risk_scores = []
        high_risk_count = 0
        risk_distribution = {'low': 0, 'medium': 0, 'high': 0, 'critical': 0}

        for doc in documents:
            try:
                analysis = json.loads(doc.get('summary', '{}'))
                score = analysis.get('riskScore', {}).get('score', 0)
                risk_scores.append(score)

                if score >= 7:
                    risk_distribution['critical'] += 1
                    high_risk_count += 1
                elif score >= 5:
                    risk_distribution['high'] += 1
                    high_risk_count += 1
                elif score >= 3:
                    risk_distribution['medium'] += 1
                else:
                    risk_distribution['low'] += 1
            except:
                pass

        avg_risk = sum(risk_scores) / len(risk_scores) if risk_scores else 0

        # Get monthly trends (last 6 months)
        monthly_trends = []
        for i in range(5, -1, -1):
            month_date = datetime.now().replace(day=1) - timedelta(days=i*30)
            month_str = month_date.strftime('%Y-%m')
            count = sync_db.documents.count_documents({
                "user_id": user_id,
                "created_at": {"$gte": month_date, "$lt": month_date + timedelta(days=30)}
            })
            monthly_trends.append({'month': month_date.strftime('%b %Y'), 'count': count})

        # Get top issues
        top_issues = []
        issue_counts = {}
        for doc in documents:
            try:
                analysis = json.loads(doc.get('summary', '{}'))
                flags = analysis.get('criticalFlags', [])
                for flag in flags:
                    title = flag.get('title', 'Unknown issue')
                    issue_counts[title] = issue_counts.get(title, 0) + 1
            except:
                pass

        top_issues = [{'issue': issue, 'count': count} for issue, count in sorted(issue_counts.items(), key=lambda x: x[1], reverse=True)[:5]]

        # Get recent activity
        recent_docs = list(sync_db.documents.find({"user_id": user_id}).sort("created_at", -1).limit(5))
        recent_activity = []
        for doc in recent_docs:
            try:
                analysis = json.loads(doc.get('summary', '{}'))
                risk_score = analysis.get('riskScore', {}).get('score', 0)
                issues_count = len(analysis.get('criticalFlags', []))
                recent_activity.append({
                    'id': doc['id'],
                    'created_at': doc.get('created_at', '').isoformat() if isinstance(doc.get('created_at'), datetime) else str(doc.get('created_at', '')),
                    'risk_score': risk_score,
                    'issues_count': issues_count
                })
            except:
                pass

        sync_client.close()

        return {
            'totalDocuments': total_docs,
            'thisMonthCount': this_month_count,
            'averageRiskScore': avg_risk,
            'highRiskCount': high_risk_count,
            'riskDistribution': risk_distribution,
            'monthlyTrends': monthly_trends,
            'topIssues': top_issues,
            'recentActivity': recent_activity
        }

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Analytics error: {e}")
        raise HTTPException(status_code=500, detail="Failed to get analytics")

@api_router.post("/compare")
async def compare_documents_endpoint(request: Request, document1: UploadFile = File(None), document2: UploadFile = File(None)):
    """Compare two documents.

    This endpoint now allows anonymous requests (no Authorization required) so users can compare two uploaded documents without signing in.
    Currently returns a simple placeholder response; file upload parameters are accepted for future processing.
    """
    try:
        # Log incoming files (if any) for debugging
        try:
            files_info = []
            if document1:
                files_info.append(f"document1={document1.filename}")
            if document2:
                files_info.append(f"document2={document2.filename}")
            logging.info(f"Compare request received. Files: {', '.join(files_info) if files_info else 'none'}")
        except Exception:
            logging.debug("No multipart files parsed or error reading file info")

        # If files were uploaded, extract text from them
        def extract_text_from_bytes(contents: bytes, filename: str, content_type: Optional[str] = None) -> str:
            text = ""
            try:
                # PDF
                if (content_type and 'pdf' in content_type) or filename.lower().endswith('.pdf'):
                    try:
                        reader = PdfReader(BytesIO(contents))
                        pages = []
                        for page in reader.pages:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                pages.append(page_text)
                        text = "\n".join(pages)
                    except Exception as e:
                        logging.error(f"PDF extraction error for {filename}: {e}")
                        text = ""
                else:
                    # Try decoding as utf-8, fall back to latin-1
                    try:
                        text = contents.decode('utf-8')
                    except Exception:
                        try:
                            text = contents.decode('latin-1')
                        except Exception:
                            text = ""
            except Exception as e:
                logging.error(f"Error extracting text from {filename}: {e}")
                text = ""

            return text or ""

        contents1 = None
        contents2 = None
        text1 = ""
        text2 = ""

        if document1:
            try:
                contents1 = await document1.read()
                text1 = extract_text_from_bytes(contents1, document1.filename, getattr(document1, 'content_type', None))
            except Exception as e:
                logging.error(f"Error reading document1: {e}")

        if document2:
            try:
                contents2 = await document2.read()
                text2 = extract_text_from_bytes(contents2, document2.filename, getattr(document2, 'content_type', None))
            except Exception as e:
                logging.error(f"Error reading document2: {e}")

        # If no uploaded files or no text extracted, return an informative error
        if not text1.strip() and not text2.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted from the uploaded documents")

        # Compute similarity
        try:
            matcher = difflib.SequenceMatcher(None, text1, text2)
            ratio = matcher.ratio()
            similarity_percent = round(ratio * 100, 1)

            # Extract matching blocks (top matches)
            matches = []
            for block in sorted(matcher.get_matching_blocks(), key=lambda b: b.size, reverse=True):
                if block.size and block.size > 40:
                    snippet = text1[block.a:block.a + block.size].strip()
                    snippet = ' '.join(snippet.split())
                    if snippet and snippet not in matches:
                        matches.append(snippet[:500])
                if len(matches) >= 5:
                    break

            # Compute simple differences by line diff
            lines1 = [l.strip() for l in text1.splitlines() if l.strip()]
            lines2 = [l.strip() for l in text2.splitlines() if l.strip()]
            raw_diff = list(difflib.ndiff(lines1, lines2))
            diffs = []
            for d in raw_diff:
                if d.startswith('- '):
                    diffs.append(d[2:].strip())
                elif d.startswith('+ '):
                    diffs.append(d[2:].strip())
                if len(diffs) >= 10:
                    break

            # Risk heuristics: higher similarity -> lower risk (simple heuristic)
            risk1 = max(1, min(10, int(round((1 - ratio) * 10))))
            risk2 = risk1

            # Recommendations based on similarity
            if ratio >= 0.85:
                recommendations = f"Documents are highly similar ({similarity_percent}% match). Verify authorship and version control; minor wording differences may be acceptable."
            elif ratio >= 0.5:
                recommendations = f"Documents show moderate similarity ({similarity_percent}%). Review sections with differences and confirm intended variations."
            else:
                recommendations = f"Documents appear substantially different ({similarity_percent}% similarity). Perform a detailed legal review for both documents."

            return {
                "similarities": matches or [f"Documents similarity: {similarity_percent}%"],
                "differences": diffs or ["No granular line-level differences could be extracted"],
                "doc1Risk": risk1,
                "doc2Risk": risk2,
                "recommendations": recommendations,
                "similarity_percent": similarity_percent
            }
        except Exception as e:
            logging.error(f"Error computing comparison: {e}")
            raise HTTPException(status_code=500, detail="Failed to compare documents")

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Compare documents error: {e}")
        raise HTTPException(status_code=500, detail="Comparison failed")

@api_router.get("/auth/me")
async def get_current_user(request: Request):
    """Get current authenticated user info"""
    try:
        # Extract token from Authorization header
        auth_header = request.headers.get("Authorization")
        if not auth_header or not auth_header.startswith("Bearer "):
            raise HTTPException(status_code=401, detail="Authorization header missing or invalid")

        token = auth_header.replace("Bearer ", "")
        payload = verify_token(token)
        if not payload:
            raise HTTPException(status_code=401, detail="Invalid token")

        user_id = payload.get("user_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token payload")

        # Get user from database - use ObjectId for _id lookup
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        try:
            user = sync_db.users.find_one({"_id": ObjectId(user_id)})
        except:
            user = sync_db.users.find_one({"id": user_id})
        sync_client.close()

        if not user:
            raise HTTPException(status_code=404, detail="User not found")

        return {
            "id": str(user.get("_id", user.get("id", ""))),
            "name": user.get("name", ""),
            "email": user.get("email", "")
        }

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Get current user error: {e}")
        raise HTTPException(status_code=500, detail="Failed to get user info")


@api_router.patch("/auth/me")
async def update_current_user(request: Request, data: dict = Body(...)):
    """Update current authenticated user's profile fields: name, description, avatar (base64 string).

    Expects Authorization: Bearer <token> header. Returns updated user object.
    """
    try:
        auth_header = request.headers.get("Authorization")
        if not auth_header or not auth_header.startswith("Bearer "):
            raise HTTPException(status_code=401, detail="Authorization header missing or invalid")

        token = auth_header.replace("Bearer ", "")
        payload = verify_token(token)
        if not payload:
            raise HTTPException(status_code=401, detail="Invalid token")

        user_id = payload.get("user_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token payload")

        update_fields = {}
        if 'name' in data and isinstance(data.get('name'), str):
            update_fields['name'] = data.get('name').strip()
        if 'description' in data and isinstance(data.get('description'), str):
            update_fields['description'] = data.get('description').strip()
        if 'avatar' in data and isinstance(data.get('avatar'), str):
            # store base64 data URL or any string the client sends
            update_fields['avatar'] = data.get('avatar')

        if not update_fields:
            raise HTTPException(status_code=400, detail="No valid fields to update")

        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        try:
            result = sync_db.users.update_one({"_id": ObjectId(user_id)}, {"$set": update_fields})
        except:
            result = sync_db.users.update_one({"id": user_id}, {"$set": update_fields})
        
        if result.matched_count == 0:
            sync_client.close()
            raise HTTPException(status_code=404, detail="User not found")

        try:
            user = sync_db.users.find_one({"_id": ObjectId(user_id)})
        except:
            user = sync_db.users.find_one({"id": user_id})
        sync_client.close()

        # Normalize returned user
        resp = {
            "id": str(user.get("_id", user.get("id", ""))), 
            "name": user.get('name'), 
            "email": user.get('email'), 
            "description": user.get('description', ''), 
            "avatar": user.get('avatar', None)
        }
        return resp

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Update current user error: {e}")
        raise HTTPException(status_code=500, detail="Failed to update user profile")

@api_router.get("/auth/google")
async def google_auth():
    """Initiate Google OAuth flow"""
    if not GOOGLE_CLIENT_ID:
        raise HTTPException(status_code=500, detail="Google OAuth not configured")

    state = secrets.token_urlsafe(32)
    google_auth_url = (
        "https://accounts.google.com/o/oauth2/v2/auth?"
        f"client_id={GOOGLE_CLIENT_ID}&"
        "response_type=code&"
        f"scope=email%20profile&"
        f"redirect_uri={GOOGLE_REDIRECT_URI}&"
        f"state={state}&"
        "access_type=offline"
    )

    return {"auth_url": google_auth_url, "state": state}

@api_router.post("/auth/google/callback", response_model=TokenResponse)
async def google_auth_callback(code: str = Form(...), state: str = Form(...)):
    """Handle Google OAuth callback"""
    try:
        if not GOOGLE_CLIENT_ID or not GOOGLE_CLIENT_SECRET:
            raise HTTPException(status_code=500, detail="Google OAuth not configured")

        # Exchange code for access token
        token_url = "https://oauth2.googleapis.com/token"
        token_data = {
            "client_id": GOOGLE_CLIENT_ID,
            "client_secret": GOOGLE_CLIENT_SECRET,
            "code": code,
            "grant_type": "authorization_code",
            "redirect_uri": GOOGLE_REDIRECT_URI,
        }

        token_response = requests.post(token_url, data=token_data)
        if token_response.status_code != 200:
            raise HTTPException(status_code=400, detail="Failed to get access token from Google")

        token_json = token_response.json()
        access_token = token_json.get("access_token")

        # Get user info from Google
        user_info_url = "https://www.googleapis.com/oauth2/v2/userinfo"
        user_info_response = requests.get(
            user_info_url,
            headers={"Authorization": f"Bearer {access_token}"}
        )

        if user_info_response.status_code != 200:
            raise HTTPException(status_code=400, detail="Failed to get user info from Google")

        user_info = user_info_response.json()
        google_id = user_info.get("id")
        email = user_info.get("email")
        name = user_info.get("name")

        # Check if user exists, create if not
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        user = sync_db.users.find_one({"$or": [{"email": email}, {"google_id": google_id}]})

        if not user:
            # Create new user
            user = User(
                email=email,
                name=name,
                google_id=google_id
            )
            user_dict = user.model_dump()
            user_dict['created_at'] = user_dict['created_at'].isoformat()
            sync_db.users.insert_one(user_dict)
        else:
            # Update existing user with Google ID if not present
            if not user.get('google_id'):
                sync_db.users.update_one(
                    {"_id": user["_id"]},
                    {"$set": {"google_id": google_id}}
                )

        sync_client.close()

        # Create JWT token
        jwt_token = create_access_token(
            data={"sub": email, "user_id": user['id'] if 'id' in user else user_dict['id']},
            expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        )

        return TokenResponse(
            access_token=jwt_token,
            user={"id": user.get('id', user_dict['id']), "email": email, "name": name}
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Google OAuth callback error: {e}")
        raise HTTPException(status_code=500, detail="OAuth authentication failed")


# Export user data (documents + chats) as a zip
@api_router.get('/auth/export')
async def export_user_data(request: Request):
    auth_header = request.headers.get('Authorization')
    token = None
    user_id = None
    if auth_header and auth_header.startswith('Bearer '):
        token = auth_header.replace('Bearer ', '')
        payload = verify_token(token)
        if payload:
            user_id = payload.get('user_id')

    # If no authenticated user, return guest history as zip if exists
    try:
        import zipfile
        buffer = BytesIO()
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            if user_id:
                sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
                sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
                docs = list(sync_db.documents.find({'user_id': user_id}))
                chats = list(sync_db.chat_messages.find({'user_id': user_id}))
                sync_client.close()

                zf.writestr('documents.json', json.dumps(docs, default=str))
                zf.writestr('chats.json', json.dumps(chats, default=str))
            else:
                # Guest: include local server-side guest documents if any
                sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
                sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
                docs = list(sync_db.documents.find({'guest': True}).limit(50)) if 'documents' in sync_db.list_collection_names() else []
                sync_client.close()
                zf.writestr('guest_documents.json', json.dumps(docs, default=str))

        buffer.seek(0)
        return StreamingResponse(buffer, media_type='application/zip', headers={"Content-Disposition": "attachment; filename=covenantai-export.zip"})
    except Exception as e:
        logging.error(f"Export failed: {e}")
        raise HTTPException(status_code=500, detail='Export failed')


# Request password reset (generates a token and logs it / emails if configured)
@api_router.post('/auth/forgot-password')
async def forgot_password(payload: dict = Body(...)):
    email = payload.get('email')
    if not email:
        raise HTTPException(status_code=400, detail='Email is required')
    sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
    sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
    user = sync_db.users.find_one({'email': email})
    if not user:
        sync_client.close()
        # Do not reveal whether the email exists
        return {'message': 'If an account exists we sent instructions to the email provided.'}

    token = secrets.token_urlsafe(20)
    expire = (datetime.utcnow() + timedelta(hours=1)).isoformat()
    sync_db.users.update_one({'email': email}, {'$set': {'pw_reset_token': token, 'pw_reset_expires': expire}})
    sync_client.close()

    # Log and optionally send email
    logging.info(f"Password reset token for {email}: {token}")
    # If SMTP configured, attempt sending (best-effort)
    try:
        smtp_host = os.environ.get('SMTP_HOST')
        if smtp_host:
            # Implement sending via smtplib or external service (omitted here)
            logging.info('Would send password reset email (SMTP configured)')
    except Exception as e:
        logging.warning('Failed to send password reset email: %s', e)

    return {'message': 'If an account exists we sent instructions to the email provided.'}


# Reset password using token
@api_router.post('/auth/reset-password')
async def reset_password(payload: dict = Body(...)):
    token = payload.get('token')
    new_password = payload.get('password')
    if not token or not new_password:
        raise HTTPException(status_code=400, detail='Token and new password are required')
    try:
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        user = sync_db.users.find_one({'pw_reset_token': token})
        if not user:
            sync_client.close()
            raise HTTPException(status_code=400, detail='Invalid or expired token')
        expires = user.get('pw_reset_expires')
        if not expires or datetime.fromisoformat(expires) < datetime.utcnow():
            sync_client.close()
            raise HTTPException(status_code=400, detail='Token expired')

        hashed = hash_password(new_password)
        sync_db.users.update_one({'_id': user['_id']}, {'$set': {'password': hashed}, '$unset': {'pw_reset_token': '', 'pw_reset_expires': ''}})
        sync_client.close()
        return {'message': 'Password reset successful'}
    except Exception as e:
        logging.error(f"Reset password error: {e}")
        raise HTTPException(status_code=500, detail='Failed to reset password')


# Email verification (generate code)
@api_router.post('/auth/send-verification')
async def send_verification(payload: dict = Body(...)):
    email = payload.get('email')
    password = payload.get('password')
    name = payload.get('name')
    
    logging.info(f"🔵 /auth/send-verification called for {email}")
    logging.info(f"   Password: {password[:5] if password else 'NONE'}...")
    logging.info(f"   Name: {name}")
    
    if not email:
        raise HTTPException(status_code=400, detail='Email required')
    try:
        # Log SMTP configuration (excluding password)
        smtp_host = os.environ.get('SMTP_HOST')
        smtp_port = os.environ.get('SMTP_PORT', '587')
        smtp_user = os.environ.get('SMTP_USER')
        smtp_pass = os.environ.get('SMTP_PASS')
        from_email = os.environ.get('SMTP_FROM')
        
        smtp_config = {
            'SMTP_HOST': smtp_host,
            'SMTP_PORT': smtp_port,
            'SMTP_USER': smtp_user,
            'SMTP_FROM': from_email,
            'SMTP_CONFIGURED': bool(smtp_host and smtp_user and smtp_pass)
        }
        logging.info(f"SMTP Configuration: {smtp_config}")
        
        code = secrets.token_hex(3)  # 6 character verification code
        logging.info(f"Generated verification code for {email}")
        
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        
        # Check if user exists
        user = sync_db.users.find_one({'email': email})
        
        if not user:
            # User doesn't exist - create them if password is provided (signup flow)
            if not password:
                logging.warning(f"Attempted to send verification code to non-existent user without password: {email}")
                sync_client.close()
                # Don't reveal that the user doesn't exist
                return {'message': 'If an account exists with this email, a verification code has been sent.'}
            
            # Create new user with verification code
            logging.info(f"🟢 Creating new user for {email}")
            hashed_pw = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
            logging.info(f"   Hashed password: {hashed_pw[:20]}...")
            new_user = {
                "email": email,
                "name": name or "",
                "password": hashed_pw,
                "email_verified": False,
                "email_verify_code": code,
                "email_verify_expires": (datetime.utcnow() + timedelta(minutes=15)).isoformat(),
                "created_at": datetime.utcnow().isoformat(),
            }
            logging.info(f"   New user dict: {new_user}")
            inserted = sync_db.users.insert_one(new_user)
            logging.info(f"✅ Created new user for email {email} with verification code: {code}")
            logging.info(f"   Inserted ID: {inserted.inserted_id}")
            logging.info(f"   Email verified: False")
            logging.info(f"   Code expires: {new_user['email_verify_expires']}")
        else:
            # Update existing user with new verification code
            sync_db.users.update_one({'email': email}, {
                '$set': {
                    'email_verify_code': code,
                    'email_verify_expires': (datetime.utcnow() + timedelta(minutes=15)).isoformat()
                }
            }, upsert=False)
        sync_client.close()
        logging.info(f"Updated verification code in database for {email}")

        # Prepare email settings
        smtp_host = os.environ.get('SMTP_HOST')
        smtp_port = int(os.environ.get('SMTP_PORT', '587'))
        smtp_user = os.environ.get('SMTP_USER')
        smtp_pass = os.environ.get('SMTP_PASS')
        from_email = os.environ.get('SMTP_FROM', smtp_user)

        if smtp_host and smtp_user and smtp_pass:
            try:
                import smtplib
                from email.mime.text import MIMEText
                from email.mime.multipart import MIMEMultipart

                message = MIMEMultipart()
                message['From'] = from_email
                message['To'] = email
                message['Subject'] = 'Your Verification Code'

                body = f"""
                Hello,

                Your verification code is: {code}

                This code will expire in 15 minutes.

                If you did not request this code, please ignore this email.

                Best regards,
                LegalDocAI Team
                """

                message.attach(MIMEText(body, 'plain'))

                logging.info(f"Connecting to SMTP server: {smtp_host}:{smtp_port}")
                server = smtplib.SMTP(smtp_host, smtp_port)
                server.set_debuglevel(1)  # Enable SMTP debug logging
                logging.info("Enabling TLS")
                server.starttls()
                logging.info("Attempting SMTP login")
                server.login(smtp_user, smtp_pass)
                logging.info("Sending message")
                server.send_message(message)
                logging.info("Message sent successfully")
                server.quit()
                logging.info("SMTP connection closed")

                logging.info(f"Verification email sent to {email}")
                # Always show the code in logs for development
                print("\n==================================")
                print(f"🔑 VERIFICATION CODE for {email}: {code}")
                print("==================================\n")
            except Exception as mail_error:
                logging.error(f"Failed to send verification email: {mail_error}")
                # Still show the code even if email fails
                print("\n==================================")
                print(f"🔑 VERIFICATION CODE for {email}: {code}")
                print("==================================\n")
        else:
            # Show code when SMTP is not configured
            print("\n==================================")
            print(f"🔑 VERIFICATION CODE for {email}: {code}")
            print("==================================\n")

        # Don't reveal whether email exists
        return {'message': 'If an account exists with this email, a verification code has been sent.'}
    except Exception as e:
        logging.error(f"Send verification error: {e}")
        raise HTTPException(status_code=500, detail='Failed to send verification')


# Verify email code
@api_router.post('/auth/verify')
async def verify_email(payload: dict = Body(...)):
    email = payload.get('email')
    code = payload.get('code')
    if not email or not code:
        raise HTTPException(status_code=400, detail='Email and code required')
    try:
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        user = sync_db.users.find_one({'email': email})
        if not user:
            sync_client.close()
            raise HTTPException(status_code=400, detail='Invalid email or code')
        
        # Normalize codes: strip whitespace and convert to lowercase for case-insensitive comparison
        stored_code = user.get('email_verify_code', '').strip().lower()
        input_code = code.strip().lower()
        
        # Check if code matches and hasn't expired
        if stored_code != input_code or datetime.fromisoformat(user.get('email_verify_expires', '1970-01-01T00:00:00')) < datetime.utcnow():
            sync_client.close()
            raise HTTPException(status_code=400, detail='Invalid or expired code')
        sync_db.users.update_one({'_id': user['_id']}, {'$set': {'email_verified': True}, '$unset': {'email_verify_code': '', 'email_verify_expires': ''}})
        sync_client.close()
        return {'message': 'Email verified'}
    except HTTPException as http_e:
        # Re-raise HTTP exceptions to preserve the status code
        raise http_e
    except Exception as e:
        logging.error(f"Verify email error: {e}")
        raise HTTPException(status_code=500, detail='Failed to verify email')


# Delete current authenticated user
@api_router.delete('/auth/me')
async def delete_my_account(request: Request):
    auth_header = request.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        raise HTTPException(status_code=401, detail='Missing authorization')
    token = auth_header.replace('Bearer ', '')
    payload = verify_token(token)
    if not payload:
        raise HTTPException(status_code=401, detail='Invalid token')
    user_id = payload.get('user_id')
    if not user_id:
        raise HTTPException(status_code=401, detail='Invalid token payload')
    try:
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        # Delete user docs and chats
        sync_db.documents.delete_many({'user_id': user_id})
        sync_db.chat_messages.delete_many({'user_id': user_id})
        sync_db.users.delete_one({'id': user_id})
        sync_client.close()
        return {'message': 'Account deleted'}
    except Exception as e:
        logging.error(f"Delete account error: {e}")
        raise HTTPException(status_code=500, detail='Failed to delete account')

# Document upload endpoint
@api_router.post("/documents/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload a legal document for analysis"""
    logging.info(f"Received upload request for file: {file.filename}")
    try:
        # Validate file type
        allowed_types = ['application/pdf', 'text/plain', 'application/msword',
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        'application/octet-stream']
        file_extension = Path(file.filename).suffix.lower()
        allowed_extensions = ['.pdf', '.doc', '.docx', '.txt']
        if file.content_type not in allowed_types and file_extension not in allowed_extensions:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.content_type} or extension: {file_extension}. Allowed types: {', '.join(allowed_types)} and extensions: {', '.join(allowed_extensions)}")
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        file_extension = Path(file.filename).suffix
        unique_filename = f"{file_id}{file_extension}"
        file_path = UPLOAD_DIR / unique_filename
        
        # Read file content fully before writing
        contents = await file.read()
        
        # Save file
        try:
            logging.info(f"Attempting to save file to: {file_path}")
            with open(file_path, "wb") as buffer:
                buffer.write(contents)
            logging.info("File saved successfully")

            # Extract text content using enhanced extraction system
            try:
                text_content, extraction_method = extract_document_text(str(file_path), file.content_type)
                
                if extraction_method == "failed" or len(text_content.strip()) < 50:
                    raise HTTPException(
                        status_code=400,
                        detail="Could not extract meaningful text from document. Please ensure the document contains readable text content."
                    )
                
                # Create enhanced document record with content stats
                content_stats = {
                    "char_count": len(text_content),
                    "word_count": len(text_content.split()),
                    "extraction_method": extraction_method,
                    "extraction_time": datetime.now(timezone.utc).isoformat()
                }
                
                document = Document(
                    filename=file.filename,
                    file_path=str(file_path),
                    file_type=file.content_type,
                    content=text_content,
                    analysis_status="pending",
                    metadata={"content_stats": content_stats}
                )
                
                # Save to database (using synchronous approach to avoid event loop conflicts)
                document_dict = document.model_dump()
                document_dict['upload_date'] = document_dict['upload_date'].isoformat()
                document_dict['content'] = text_content  # Store the extracted text
                document_dict['analysis_status'] = 'pending'

                # Use synchronous MongoDB client for this operation
                import pymongo
                sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
                sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
                sync_db.documents.insert_one(document_dict)
                
                # Trigger document analysis in the background
                asyncio.create_task(analyze_document(document.id))
                
                sync_client.close()
                
                return {"document_id": document.id, "filename": document.filename, "status": "uploaded"}
            
            except Exception as e:
                logging.error(f"Error extracting text or saving to database: {str(e)}")
                # Clean up the file if it was created but there was an error
                if file_path.exists():
                    try:
                        file_path.unlink()
                    except:
                        pass
                raise HTTPException(status_code=500, detail="Document processing failed")
            
        except Exception as e:
            logging.error(f"Error saving file: {str(e)}")
            # Clean up the file if it was created but there was a database error
            if file_path.exists():
                try:
                    file_path.unlink()
                except:
                    pass
            raise
        
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Upload error: {str(e)}")
        raise HTTPException(status_code=500, detail="File upload failed")

# Get all documents
@api_router.get("/documents", response_model=List[Document])
async def get_documents():
    """Get all uploaded documents"""
    try:
        # Use synchronous MongoDB client to avoid event loop conflicts
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        documents = list(sync_db.documents.find())
        sync_client.close()

        safe_docs = []
        for doc in documents:
            # Normalize required fields with safe defaults to avoid Pydantic validation errors
            safe = {}
            # id: prefer 'id' field, otherwise fallback to MongoDB _id string
            safe['id'] = doc.get('id') or str(doc.get('_id'))
            safe['filename'] = doc.get('filename', 'Untitled Document')
            safe['file_path'] = doc.get('file_path', '')
            safe['file_type'] = doc.get('file_type', 'application/octet-stream')

            # upload_date may be stored as datetime or ISO string; normalize to datetime
            ud = doc.get('upload_date')
            if isinstance(ud, datetime):
                safe['upload_date'] = ud
            elif isinstance(ud, str):
                try:
                    safe['upload_date'] = datetime.fromisoformat(ud)
                except Exception:
                    safe['upload_date'] = datetime.now(timezone.utc)
            else:
                safe['upload_date'] = datetime.now(timezone.utc)

            safe['analysis_status'] = doc.get('analysis_status', 'completed')
            safe['summary'] = doc.get('summary')
            safe['key_clauses'] = doc.get('key_clauses')
            safe['risk_assessment'] = doc.get('risk_assessment')

            safe_docs.append(safe)

        return [Document(**d) for d in safe_docs]
    except Exception as e:
        logging.error(f"Get documents error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve documents")

# Delete document endpoint
@api_router.delete("/documents/{document_id}")
async def delete_document(document_id: str):
    """Delete a document by ID"""
    try:
        # Log the delete request
        logging.info(f"Attempting to delete document with ID: {document_id}")

        # Use synchronous MongoDB client to avoid event loop conflicts
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        # Try to find by 'id' field first, then by '_id' field (MongoDB ObjectId)
        document = sync_db.documents.find_one({"id": document_id})
        if not document:
            try:
                # Try as MongoDB ObjectId
                from bson import ObjectId
                document = sync_db.documents.find_one({"_id": ObjectId(document_id)})
            except Exception:
                pass
        
        if not document:
            sync_client.close()
            logging.warning(f"Document not found for deletion: {document_id}")
            raise HTTPException(status_code=404, detail="Document not found")

        # Delete the file from the filesystem
        try:
            file_path = Path(document.get('file_path', ''))
            if file_path and file_path.exists():
                file_path.unlink()
                logging.info(f"File deleted successfully: {file_path}")
        except Exception as e:
            logging.error(f"Error deleting file: {e}")
            # Continue with document deletion even if file deletion fails

        # Delete the document record from the database - use the actual _id
        result = sync_db.documents.delete_one({"_id": document["_id"]})
        sync_client.close()

        if result.deleted_count == 0:
            logging.warning(f"Document not deleted from database: {document_id}")
            raise HTTPException(status_code=404, detail="Document not found in database")

        logging.info(f"Document deleted successfully: {document_id}")
        return {"message": "Document deleted successfully", "id": document_id}

    except HTTPException as he:
        raise he
    except Exception as e:
        logging.error(f"Delete document error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Get specific document
@api_router.get("/documents/{document_id}")
async def get_document(document_id: str):
    """Get a specific document by ID"""
    try:
        # Use synchronous MongoDB client to avoid event loop conflicts
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        document = sync_db.documents.find_one({"id": document_id})
        sync_client.close()

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        if isinstance(document.get('upload_date'), str):
            document['upload_date'] = datetime.fromisoformat(document['upload_date'])

        return Document(**document)
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Get document error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve document")




MAX_CHUNK_SIZE = 8000

def generate_analysis_prompt(text: str) -> str:
    return (
        "You are a legal document analysis expert. Analyze this document content and provide a "
        "structured analysis in JSON format.\n\n"
        f"DOCUMENT CONTENT:\n{text}\n\n"
        "Provide your response as a JSON object with these exact fields:\n"
        "{\n"
        '  "document_type": {"type": "string", "confidence": "High/Medium/Low"},\n'
        '  "key_points": [\n'
        '    {"point": "string", "importance": "Critical/High/Medium/Low"}\n'
        "  ],\n"
        '  "parties": [\n'
        '    {"name": "string", "role": "string", "obligations": ["string"]}\n'
        "  ],\n"
        '  "risks": [\n'
        '    {"risk": "string", "severity": "High/Medium/Low", "mitigation": "string"}\n'
        "  ],\n"
        '  "recommendations": [\n'
        '    {"action": "string", "priority": "High/Medium/Low", "rationale": "string"}\n'
        "  ]\n"
        "}\n\n"
        "REQUIREMENTS:\n"
        "1. Base analysis ONLY on the provided content\n"
        "2. Include specific references to document text\n"
        "3. Focus on legal implications and risks\n"
        "4. Keep explanations clear and actionable\n"
        "5. Ensure output is valid JSON"
    )

# Placeholder extraction helpers - replace with real implementations
def extract_text_pdf_fallback(path: str) -> str:
    return ""

def extract_text_with_ocr(path: str) -> str:
    return ""

def extract_text_docx(path: str) -> str:
    return ""

@api_router.post("/documents/{document_id}/analyze")
async def analyze_document(document_id: str, request: Request):
    sync_client = None
    user_id = None
    all_analyses: List[Dict[str, Any]] = []
    try:
        # Extract user_id from Authorization header if present
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.startswith("Bearer "):
            token = auth_header.replace("Bearer ", "")
            payload = verify_token(token)
            if payload:
                user_id = payload.get("user_id")
                logging.info(f"Analyze request from user: {user_id}")
        
        # Connect to DB
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]

        document = sync_db.documents.find_one({"id": document_id})
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # Mark processing
        sync_db.documents.update_one({"id": document_id}, {"$set": {"analysis_status": "processing"}})

        # Get content or fallback to reading file
        content = document.get('content', '') or document.get('document_text', '')
        document_text = content or ""

        # If no inline content, try to read from file_path
        if not document_text or len(document_text.strip()) < 50:
            file_path = document.get('file_path')
            if file_path:
                fp = Path(file_path)
                if not fp.exists():
                    logging.warning(f"File not found at path: {file_path}")
                else:
                    # If PDF, use PyPDF2 then fallbacks
                    try:
                        if document.get('file_type') == 'application/pdf' or fp.suffix.lower() == '.pdf':
                            from PyPDF2 import PdfReader
                            reader = PdfReader(str(fp))
                            pages = []
                            for p in reader.pages:
                                txt = p.extract_text() or ''
                                pages.append(txt)
                            document_text = "\n".join(pages).strip()
                            if len(document_text) < 50:
                                pdf_text = extract_text_pdf_fallback(str(fp))
                                if pdf_text and len(pdf_text.strip()) > len(document_text):
                                    document_text = pdf_text
                                else:
                                    ocr_text = extract_text_with_ocr(str(fp))
                                    if ocr_text and len(ocr_text.strip()) > len(document_text):
                                        document_text = ocr_text
                        else:
                            # docx fallback or plain text
                            if fp.suffix.lower() == '.docx':
                                document_text = extract_text_docx(str(fp)) or document_text
                            if not document_text:
                                # try reading as utf-8
                                with open(fp, 'r', encoding='utf-8') as f:
                                    document_text = f.read()
                    except Exception as e:
                        logging.exception("Error extracting text from file")
                        # keep document_text as-is or empty

        if not document_text or len(document_text.strip()) < 50:
            sync_db.documents.update_one({"id": document_id}, {"$set": {"analysis_status": "failed", "analysis_error": "Insufficient text for analysis"}})
            raise HTTPException(status_code=400, detail="Document has insufficient content for analysis")

        logging.info(f"Starting analysis for document {document_id}, extracted length={len(document_text)}")

        # Chunk and analyze
        chunks = [document_text[i:i + MAX_CHUNK_SIZE] for i in range(0, len(document_text), MAX_CHUNK_SIZE)]
        for i, chunk in enumerate(chunks):
            logging.info(f"Analyzing chunk {i+1}/{len(chunks)}")
            prompt = generate_analysis_prompt(chunk)
            try:
                resp = await send_message(prompt)
                resp = resp.strip()
                # remove triple-backtick fences if present
                if resp.startswith(""):
                    # support json or 
                    resp = resp.split("\n", 1)[1] if "\n" in resp else resp[3:]
                    if resp.endswith(""):
                        resp = resp[:-3]
                    resp = resp.strip()
                parsed = json.loads(resp)
                all_analyses.append(parsed)
            except json.JSONDecodeError:
                logging.exception("AI returned non-JSON for a chunk; skipping that chunk")
                continue
            except Exception:
                logging.exception("Error calling send_message for chunk")
                continue

        if not all_analyses:
            sync_db.documents.update_one({"id": document_id}, {"$set": {"analysis_status": "failed", "analysis_error": "AI returned no valid JSON responses"}})
            raise HTTPException(status_code=500, detail="Failed to get valid analysis from AI")

        # Merge analyses
        merged = {
            "document_type": all_analyses[0].get("document_type", {}),
            "key_points": [],
            "risks": [],
            "recommendations": []
        }
        seen_points = set()
        seen_risks = set()
        seen_recs = set()

        for analysis in all_analyses:
            for kp in analysis.get("key_points", []):
                text = kp.get("point", "").strip()
                if text and text not in seen_points:
                    merged["key_points"].append(kp)
                    seen_points.add(text)
            for r in analysis.get("risks", []):
                text = r.get("risk", "").strip()
                if text and text not in seen_risks:
                    merged["risks"].append(r)
                    seen_risks.add(text)
            for rc in analysis.get("recommendations", []):
                text = rc.get("action", "").strip()
                if text and text not in seen_recs:
                    merged["recommendations"].append(rc)
                    seen_recs.add(text)

        # Prepare update dict
        update_dict = {
            "analysis_status": "completed",
            "document_type": merged["document_type"],
            "key_points": merged["key_points"],
            "risks": merged["risks"],
            "recommendations": merged["recommendations"],
            "analysis_version": "2.0",
            "last_analyzed": datetime.now(timezone.utc).isoformat()
        }
        
        # If user is authenticated, add user_id to mark this as user's document
        if user_id:
            update_dict["user_id"] = user_id
            logging.info(f"✅ Saving analysis with user_id: {user_id}")
        
        # Update DB with completed analysis
        sync_db.documents.update_one(
            {"id": document_id},
            {"$set": update_dict}
        )

        logging.info(f"Analysis completed for document {document_id}")
        return merged

    except HTTPException:
        # re-raise to let FastAPI return it unchanged
        raise
    except Exception as e:
        logging.exception(f"Unexpected error analyzing document {document_id}")
        # Attempt to mark failed
        try:
            if sync_client:
                sync_db.documents.update_one({"id": document_id}, {"$set": {"analysis_status": "failed", "analysis_error": str(e)}})
        except Exception:
            logging.exception("Failed to write failure status to DB")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {e}")
    finally:
        if sync_client:
            try:
                sync_client.close()
            except Exception:
                logging.exception("Error closing Mongo client")

# New /api/analyze endpoint for direct text analysis and file upload
@api_router.post("/analyze")
async def analyze_text(file: UploadFile = File(None), rawText: str = Form(None)):
    """Analyze raw text or uploaded file and return structured JSON response"""
    try:
        document_text = ""

        if file:
            logging.info(f"Received file upload analysis request. Filename: {file.filename}, Content-Type: {file.content_type}")

            # Validate file type
            allowed_types = ['application/pdf', 'text/plain', 'application/msword',
                            'application/vnd.openxmlformats-officedocument.wordprocessingml.document']
            if file.content_type not in allowed_types and not file.filename.lower().endswith(('.pdf', '.doc', '.docx', '.txt')):
                raise HTTPException(status_code=400, detail=f"Unsupported file type: {file.content_type}")

            # Read file content
            contents = await file.read()

            if file.content_type == 'application/pdf' or file.filename.lower().endswith('.pdf'):
                # Extract text from PDF
                try:
                    from io import BytesIO
                    pdf_reader = PdfReader(BytesIO(contents))
                    text_pages = []
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_pages.append(page_text)
                    document_text = "\n".join(text_pages)

                    if len(document_text.strip()) < 50:
                        document_text = "This PDF appears to be image-based or scanned. OCR processing is not available. Please provide a text-based PDF or paste the document text directly."
                except Exception as e:
                    document_text = f"Failed to extract text from PDF: {str(e)}"
            else:
                # For text files
                try:
                    document_text = contents.decode('utf-8')
                except UnicodeDecodeError:
                    document_text = "Unable to decode file content. Please ensure it's a text-based document."

        if rawText:
            logging.info(f"Received direct text analysis request. Text length: {len(rawText)}")
            document_text = rawText

        if not document_text or len(document_text.strip()) < 10:
            raise HTTPException(status_code=400, detail="Document content is too short for analysis")

        # Create analysis prompt for the document text
        analysis_prompt = f"""You are a legal document analysis expert. Analyze the following legal document text and provide a comprehensive analysis in the exact JSON format specified below.

LEGAL DOCUMENT TEXT TO ANALYZE:
{document_text[:120000]}

IMPORTANT: Provide your response as a valid JSON object with these exact keys and structure:

{{
  "summary": "A clear, concise summary of what this document is about in plain English",
  "riskScore": {{
    "score": 7,
    "max": 10,
    "label": "High Risk - Requires Immediate Legal Review"
  }},
  "analysis": {{
    "strengths": [
      {{"text": "Clear and unambiguous language"}},
      {{"text": "Well-defined obligations for both parties"}},
      {{"text": "Reasonable termination clauses"}}
    ],
    "weaknesses": [
      {{"text": "Unfavorable payment terms"}},
      {{"text": "Limited liability protection"}},
      {{"text": "Ambiguous dispute resolution"}}
    ],
    "opportunities": [
      {{"text": "Negotiate better payment terms"}},
      {{"text": "Add stronger liability protections"}},
      {{"text": "Clarify dispute resolution process"}}
    ],
    "threats": [
      {{"text": "Potential financial losses"}},
      {{"text": "Legal disputes and litigation costs"}},
      {{"text": "Damage to business reputation"}}
    ]
  }},
  "criticalFlags": [
    {{
      "title": "Unfavorable Payment Terms",
      "explanation": "The payment terms heavily favor the other party with delayed payments and high penalties.",
      "source": "Section 4.2 Payment Terms"
    }},
    {{
      "title": "Weak Liability Protection",
      "explanation": "Liability limitations are insufficient and may not hold up in court.",
      "source": "Section 7.1 Limitation of Liability"
    }}
  ],
  "negotiationPoints": [
    {{
      "title": "Payment Terms",
      "risk": "High - Could impact cash flow significantly",
      "example": "Request 50% payment upfront and 50% upon completion"
    }},
    {{
      "title": "Liability Cap",
      "risk": "Medium - May expose to unlimited liability",
      "example": "Cap liability at 2x the contract value"
    }}
  ]
}}

REQUIREMENTS:
- Base ALL analysis on the provided document text
- Use the exact JSON structure shown above
- Provide realistic, specific analysis based on the document content
- Include 3-5 items in each analysis array
- Make criticalFlags and negotiationPoints specific to the document
- Ensure the response is valid JSON that can be parsed

If the document is too short or unclear, provide a general analysis framework but still use the required JSON structure."""

        # Get AI analysis
        analysis_response = await send_message(analysis_prompt)
        logging.info(f"AI analysis response received. Length: {len(analysis_response)}")

        # Parse the JSON response
        try:
            # Clean the response to ensure it's valid JSON
            analysis_response = analysis_response.strip()
            if analysis_response.startswith('```json'):
                analysis_response = analysis_response[7:]
            if analysis_response.endswith('```'):
                analysis_response = analysis_response[:-3]
            analysis_response = analysis_response.strip()

            # Try to parse JSON
            try:
                result = json.loads(analysis_response)
                logging.info("Successfully parsed analysis JSON")
            except json.JSONDecodeError as e:
                logging.error(f"JSON parsing error: {e}")
                logging.error(f"Raw response: {analysis_response[:1000]}")

                # Try to extract JSON from the response if it's wrapped in other text
                import re
                json_match = re.search(r'\{.*\}', analysis_response, re.DOTALL)
                if json_match:
                    try:
                        result = json.loads(json_match.group())
                        logging.info("Successfully extracted and parsed JSON from response")
                    except json.JSONDecodeError:
                        logging.warning("Failed to parse extracted JSON; will fallback to raw AI text instead of returning 500")
                        # Fall back to a safe result containing the raw AI response so downstream
                        # consumers (frontend/export) can still show the full analysis text.
                        result = {
                            "summary": (analysis_response[:200] + '...') if len(analysis_response) > 200 else analysis_response,
                            "riskScore": {"score": 0, "max": 10, "label": "Unknown"},
                            "analysis": {"raw_text": analysis_response},
                            "criticalFlags": [],
                            "negotiationPoints": []
                        }
                else:
                    logging.warning("AI response did not contain JSON; falling back to raw AI text")
                    result = {
                        "summary": (analysis_response[:200] + '...') if len(analysis_response) > 200 else analysis_response,
                        "riskScore": {"score": 0, "max": 10, "label": "Unknown"},
                        "analysis": {"raw_text": analysis_response},
                        "criticalFlags": [],
                        "negotiationPoints": []
                    }

            # Ensure required fields exist in the result (if we created a fallback above,
            # they will already be present). If some are missing, add safe defaults so
            # downstream code can rely on the structure.
            required_fields = ['summary', 'riskScore', 'analysis', 'criticalFlags', 'negotiationPoints']
            for field in required_fields:
                if field not in result:
                    logging.warning(f"Missing field '{field}' in AI result; inserting default value")
                    if field == 'summary':
                        result['summary'] = ''
                    elif field == 'riskScore':
                        result['riskScore'] = {"score": 0, "max": 10, "label": "Unknown"}
                    elif field == 'analysis':
                        result['analysis'] = {"raw_text": analysis_response}
                    else:
                        result[field] = []

            # Ensure SWOT keys exist and are lists so the frontend can safely render them
            analysis_obj = result.get('analysis', {}) if isinstance(result.get('analysis', {}), dict) else {}
            for k in ('strengths', 'weaknesses', 'opportunities', 'threats'):
                if k not in analysis_obj or not isinstance(analysis_obj.get(k), list):
                    analysis_obj[k] = []
            result['analysis'] = analysis_obj

            return result

        except json.JSONDecodeError as e:
            logging.error(f"JSON parsing error: {e}")
            logging.error(f"Raw response: {analysis_response[:500]}")
            raise HTTPException(status_code=500, detail=f"Failed to parse AI response as JSON: {str(e)}")

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Analysis error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")




@api_router.post("/documents/ask")
async def ask_question(request: QuestionRequest):
    """Ask a specific question about a document."""
    sync_client = None
    try:
        # --- Connect to database ---
        mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
        db_name = os.environ.get('DB_NAME', 'legal_docs')
        sync_client = pymongo.MongoClient(mongo_url)
        sync_db = sync_client[db_name]

        # --- Get document ---
        document = sync_db.documents.find_one({"id": request.document_id})
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        # --- Prepare session ID ---
        session_id = request.session_id or f"qa_{request.document_id}_{uuid.uuid4()}"

        # --- Prepare file content wrapper ---
        file_content = FileContentWithMimeType(
            file_path=document['file_path'],
            mime_type=document['file_type']
        )

        # --- Extract text from document ---
        document_text = ""
        try:
            if document['file_type'] == 'application/pdf' or document['file_path'].lower().endswith('.pdf'):
                # Try PDF text extraction
                try:
                    reader = PdfReader(document['file_path'])
                    text_pages = [page.extract_text() or "" for page in reader.pages]
                    document_text = "\n".join([t.strip() for t in text_pages if t and t.strip()])

                    # If extracted text too short → attempt OCR
                    if len(document_text.strip()) < 50:
                        logging.info("Attempting OCR for image-based/scanned PDF")
                        ocr_text = extract_text_with_ocr(document['file_path'])
                        if len(ocr_text.strip()) > len(document_text.strip()):
                            document_text = ocr_text
                            logging.info(f"OCR improved extraction length: {len(document_text)} chars")
                        else:
                            msg = (
                                f"This PDF seems scanned or image-based. "
                                f"Text extraction yielded only {len(document_text)} chars. "
                                "OCR attempted but results insufficient. Please use a text-based PDF."
                            )
                            logging.warning(msg)
                            document_text = msg
                    else:
                        logging.info(f"Extracted {len(document_text)} chars from PDF")
                except Exception as e:
                    logging.error(f"PDF extraction error: {e}")
                    document_text = f"Failed to extract text from PDF: {str(e)}"
            else:
                # Non-PDF files
                with open(document['file_path'], 'r', encoding='utf-8') as f:
                    document_text = f.read()
                logging.info(f"Extracted {len(document_text)} chars from text file")
        except UnicodeDecodeError:
            # Retry with bytes
            try:
                with open(document['file_path'], 'rb') as f:
                    document_bytes = f.read()
                document_text = document_bytes.decode('utf-8', errors='ignore')
            except Exception:
                document_text = "Unable to decode file; likely binary format."
        except Exception as e:
            document_text = f"Error reading document: {str(e)}"

        # --- Create question prompt ---
        question_prompt = f"""
You are a legal document analysis expert.
Below is the content of a legal document. Please answer the question based ONLY on this document.

DOCUMENT CONTENT:
{document_text[:10000]}

QUESTION: {request.question}

REQUIREMENTS:
1. Base your answer solely on this document.
2. Reference specific clauses or text segments when possible.
3. Avoid generic legal advice—cite this document’s content.
4. Provide:
   - A direct answer,
   - Supporting sections,
   - Context or implications,
   - Practical advice relevant to this document.
Keep the explanation clear but accurate to the text.
        """.strip()

        # --- Send to model ---
        question_message = UserMessage(text=question_prompt)
        answer = await send_message(question_prompt)

        if not answer or len(answer.strip()) < 10:
            raise Exception("Empty or invalid response from model")

        # --- Normalize answer ---
        answer = strip_markdown(answer)
        words = answer.split()
        if len(words) > 100:
            answer = ' '.join(words[:100]) + '...'
        answer = re.sub(r'\s+', ' ', answer).strip()

        # --- Save chat message ---
        chat_message = ChatMessage(
            document_id=request.document_id,
            session_id=session_id,
            question=request.question,
            answer=answer
        )
        chat_dict = chat_message.model_dump()
        chat_dict['timestamp'] = chat_dict['timestamp'].isoformat()

        sync_db.chat_messages.insert_one(chat_dict)

        # --- Return response ---
        return {
            "question": request.question,
            "answer": answer,
            "session_id": session_id
        }

    except HTTPException:
        raise
    except Exception as e:
        logging.exception("Error in /documents/ask")
        raise HTTPException(status_code=500, detail=f"Failed to answer question: {str(e)}")
    finally:
        if sync_client:
            try:
                sync_client.close()
            except Exception:
                logging.error("Error closing MongoDB connection")


@api_router.post("/documents/ask/inline")
async def ask_inline(payload: dict = Body(...)):
    """Ask a specific question about a provided document text or analysis (no DB entry required).

    Expected payload:
    {
      "question": "...",
      "document_text": "..."  # optional if analysis provided
      "document_id": "..." # optional, will try to load from database
      "analysis": {...} # optional, used when provided (analysis may contain raw_text)
    }
    """
    try:
        question = payload.get('question')
        if not question:
            raise HTTPException(status_code=400, detail="Question is required")

        # Prefer raw text from analysis if provided
        document_text = ''
        if payload.get('analysis') and isinstance(payload.get('analysis'), dict):
            analysis = payload.get('analysis')
            document_text = analysis.get('raw_text') or analysis.get('document_text') or ''

        # Fallback to explicit document_text
        if not document_text:
            document_text = payload.get('document_text', '')

        # If still empty, try to load from database if document_id is provided
        if not document_text or len(document_text.strip()) < 20:
            document_id = payload.get('document_id')
            logging.info(f"[INLINE ASK] document_id={document_id}, document_text_len={len(document_text.strip())}")
            if document_id:
                try:
                    sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
                    sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
                    # Try to find by 'id' field first, then by '_id' field (MongoDB ObjectId)
                    document = sync_db.documents.find_one({"id": document_id})
                    logging.info(f"[INLINE ASK] Search by id: {document is not None}")
                    if not document:
                        # Try finding by _id (in case document_id is actually a MongoDB _id)
                        try:
                            document = sync_db.documents.find_one({"_id": ObjectId(document_id)})
                            logging.info(f"[INLINE ASK] Search by _id: {document is not None}")
                        except Exception as e:
                            # If it's not a valid ObjectId, just skip this search
                            logging.debug(f"Document ID {document_id} is not a valid ObjectId: {e}")
                    sync_client.close()
                    
                    if document:
                        logging.info(f"[INLINE ASK] Document found, ID={document.get('id', 'NO_ID')}")
                        # Try to get content from database fields
                        db_text = (
                            document.get('content', '') or 
                            document.get('document_text', '') or
                            document.get('summary', '')
                        )
                        logging.info(f"[INLINE ASK] DB content length: {len(db_text) if db_text else 0}")
                        
                        # If still empty, try to read from file
                        if not db_text:
                            file_path = document.get('file_path')
                            logging.info(f"[INLINE ASK] No DB content, trying file_path={file_path}")
                            if file_path:
                                try:
                                    from pathlib import Path
                                    if Path(file_path).exists():
                                        logging.info(f"[INLINE ASK] File exists: {file_path}")
                                        if file_path.lower().endswith('.pdf'):
                                            from PyPDF2 import PdfReader
                                            reader = PdfReader(file_path)
                                            pages = []
                                            for page in reader.pages:
                                                text = page.extract_text()
                                                if text:
                                                    pages.append(text)
                                            db_text = "\n".join(pages)
                                        else:
                                            with open(file_path, 'r', encoding='utf-8') as f:
                                                db_text = f.read()
                                        logging.info(f"[INLINE ASK] Loaded from file, length={len(db_text)}")
                                    else:
                                        logging.warning(f"[INLINE ASK] File does not exist: {file_path}")
                                except Exception as e:
                                    logging.error(f"Failed to read file from inline endpoint: {e}")
                        
                        if db_text:
                            document_text = db_text
                            logging.info(f"[INLINE ASK] Loaded document content, final_length={len(document_text)}")
                    else:
                        logging.info(f"[INLINE ASK] No document found with ID {document_id}")
                except Exception as e:
                    logging.error(f"Failed to load document from database in inline endpoint: {e}")

        # If we have very little context, still call the general chat endpoint
        if not document_text or len(document_text.strip()) < 20:
            # Build a short prompt describing that document context is limited
            prompt = f"You are an expert legal assistant. The user asks: {question}. Note: limited or no document content was provided. Answer concisely."
        else:
            # Build a prompt that instructs the LLM to answer based ONLY on the provided document text
            prompt = f"You are a legal document analysis expert. Use ONLY the document text provided below to answer the question. DOCUMENT TEXT:\n{document_text[:12000]}\nQUESTION: {question}\nPlease answer directly, reference the document where appropriate, and keep the answer concise."

        answer = await send_message(prompt)
        if not answer or len(answer.strip()) < 1:
            raise Exception("Empty response from AI service")
        # Strip markdown from AI answer and optionally persist the QA pair
        answer = strip_markdown(answer)
        # Normalize and truncate similar to general chat
        words = answer.split()
        if len(words) > 100:
            answer = ' '.join(words[:100]) + '...'
        answer = re.sub(r'\s+', ' ', answer).strip()

        # Optionally persist the QA pair to chat_messages for continuity (non-blocking)
        try:
            chat_message = ChatMessage(
                document_id=payload.get('document_id', f"inline_{uuid.uuid4()}"),
                session_id=payload.get('session_id', f"inline_{uuid.uuid4()}"),
                question=question,
                answer=answer
            )
            chat_dict = chat_message.model_dump()
            chat_dict['timestamp'] = chat_dict['timestamp'].isoformat()
            sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
            sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
            sync_db.chat_messages.insert_one(chat_dict)
            sync_client.close()
        except Exception:
            # Non-fatal if saving fails
            logging.debug("Failed to persist inline chat message (non-fatal)")

        return {"answer": answer}

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Inline question answering error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to answer inline question: {str(e)}")

# Get chat history for a document
@api_router.get("/documents/{document_id}/chat")
async def get_chat_history(document_id: str, session_id: Optional[str] = None):
    """Get chat history for a document"""
    try:
        query = {"document_id": document_id}
        if session_id:
            query["session_id"] = session_id

        # Use synchronous MongoDB client to avoid event loop conflicts
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        messages = list(sync_db.chat_messages.find(query).sort("timestamp", 1))
        sync_client.close()

        for msg in messages:
            if isinstance(msg.get('timestamp'), str):
                msg['timestamp'] = datetime.fromisoformat(msg['timestamp'])

        return [ChatMessage(**msg) for msg in messages]

    except Exception as e:
        logging.error(f"Get chat history error: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve chat history")

# Export document analysis as selective PDF
@api_router.post("/documents/{document_id}/export-selective-pdf")
async def export_selective_pdf(document_id: str, request_body: ExportRequest):
    """Export selected sections of document analysis as PDF"""
    logging.info(f"Starting selective PDF export for document: {document_id}")
    try:
        sections = request_body.sections
        if not sections:
            raise HTTPException(status_code=400, detail="No sections selected for export")

        # Get document from database
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        document = sync_db.documents.find_one({"id": document_id})
        if not document:
            sync_client.close()
            logging.error(f"Document not found: {document_id}")
            raise HTTPException(status_code=404, detail="Document not found")

        if document.get('analysis_status') != 'completed':
            sync_client.close()
            logging.error(f"Document analysis not completed: {document_id}")
            raise HTTPException(status_code=400, detail="Document analysis not completed yet")

        sync_client.close()

        # Helper function to clean text for PDF
        def clean_text_for_pdf(text):
            """Minimal, unicode-preserving cleaning for PDF output.

            Keep Unicode characters (don't filter by byte value). Only remove
            nulls and normalize common punctuation so the PDF looks consistent.
            """
            if not text:
                return ""
            text = str(text)
            # Remove NULs which can break some PDF writers
            text = text.replace('\x00', '')
            # Replace bullet points with asterisks
            text = text.replace('•', '*').replace('\u2022', '*')
            # Replace en/em dashes with hyphens
            text = text.replace('–', '-').replace('—', '-')
            # Normalize smart quotes to ASCII equivalents for consistency
            text = text.replace('\u2018', "'").replace('\u2019', "'")
            text = text.replace('\u201C', '"').replace('\u201D', '"')
            return text

        logging.info(f"Generating selective PDF with sections: {sections}")
        from fpdf import FPDF

        pdf = FPDF()
        # If DejaVu Sans is available on the system, register it for Unicode support
        try:
            dejavu_path = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
            if os.path.exists(dejavu_path):
                pdf.add_font('DejaVu', '', dejavu_path, uni=True)
                # Register a bold variant (reuse same TTF if a separate bold file isn't available)
                try:
                    pdf.add_font('DejaVu', 'B', dejavu_path, uni=True)
                except Exception:
                    pass
                base_font = 'DejaVu'
            else:
                base_font = 'Arial'
        except Exception:
            base_font = 'Arial'
        logging.info(f"PDF selective export: selected base_font={base_font}")

        pdf.add_page()
        pdf.set_font(base_font, 'B', 16)
        pdf.cell(0, 10, txt="Legal Document Analysis Report", ln=True, align='C')
        pdf.ln(5)

        pdf.set_font(base_font, size=11)
        pdf.cell(0, 8, txt=f"Document: {document.get('filename', 'Unknown')}", ln=True)
        pdf.cell(0, 8, txt=f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.ln(5)

        # Add selected sections
        if 'summary' in sections:
            executive_summary = document.get('executive_summary') or ''
            plain_english = document.get('plain_english') or ''
            if executive_summary or plain_english:
                pdf.set_font(base_font, 'B', 12)
                pdf.cell(0, 8, txt="Summary", ln=True)
                pdf.set_font(base_font, size=10)
                
                if executive_summary:
                    clean_content = clean_text_for_pdf(executive_summary)
                    logging.info(f"selective export executive_summary length={len(executive_summary)} cleaned={len(clean_content)}")
                    if len(clean_content) > 3000:
                        clean_content = clean_content[:3000] + "...[Content truncated]"
                    pdf.multi_cell(0, 4, txt=clean_content)
                
                if plain_english:
                    clean_content = clean_text_for_pdf(plain_english)
                    logging.info(f"selective export plain_english length={len(plain_english)} cleaned={len(clean_content)}")
                    # Allow the full plain English summary to be written; multi_cell will
                    # handle paging. Avoid artificial truncation so the exported report
                    # contains the full analysis.
                    pdf.multi_cell(0, 4, txt=clean_content)
                pdf.ln(3)

        if 'keyclauses' in sections:
            key_clauses = document.get('key_clauses') or []
            if key_clauses:
                pdf.set_font(base_font, 'B', 12)
                pdf.cell(0, 8, txt="Key Clauses Analysis", ln=True)
                pdf.set_font(base_font, size=10)
                
                for clause in key_clauses:
                    if clause.get('clause'):
                        clean_clause = clean_text_for_pdf(clause['clause'])
                        pdf.set_font(base_font, 'B', 10)
                        pdf.cell(0, 6, txt=f"Clause: {clean_clause[:100]}", ln=True)
                        
                        if clause.get('explanation'):
                            pdf.set_font(base_font, size=9)
                            clean_explanation = clean_text_for_pdf(clause['explanation'])
                            pdf.multi_cell(0, 4, txt=clean_explanation)
                        pdf.ln(2)

        if 'risks' in sections:
            risk_assessment = document.get('risk_assessment') or ''
            if risk_assessment:
                pdf.set_font(base_font, 'B', 12)
                pdf.cell(0, 8, txt="Risk Assessment", ln=True)
                pdf.set_font(base_font, size=10)
                
                clean_content = clean_text_for_pdf(risk_assessment)
                pdf.multi_cell(0, 4, txt=clean_content)
                pdf.ln(3)

        if 'qa' in sections:
            # Get Q&A from chat messages
            sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
            sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
            qa_messages = list(sync_db.chat_messages.find({"document_id": document_id}).limit(10))
            sync_client.close()
            
            if qa_messages:
                pdf.set_font(base_font, 'B', 12)
                pdf.cell(0, 8, txt="Questions & Answers", ln=True)
                pdf.set_font(base_font, size=10)
                
                for qa in qa_messages:
                    clean_q = clean_text_for_pdf(qa.get('question', ''))
                    pdf.set_font(base_font, 'B', 10)
                    pdf.cell(0, 6, txt=f"Q: {clean_q[:100]}", ln=True)
                    
                    pdf.set_font(base_font, size=9)
                    clean_a = clean_text_for_pdf(qa.get('answer', ''))
                    pdf.multi_cell(0, 4, txt=f"A: {clean_a}")
                    pdf.ln(2)

        # Save PDF to temporary file
        logging.info("Saving selective PDF to temporary file")
        tmp_file_path = tempfile.mktemp(suffix='.pdf')
        pdf.output(tmp_file_path)
        
        logging.info(f"Selective PDF saved successfully to: {tmp_file_path}")

        # Return PDF file
        filename = f"legal_analysis_{document_id}.pdf"
        logging.info(f"Returning selective PDF file: {filename}")
        
        return FileResponse(
            tmp_file_path,
            media_type='application/pdf',
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Selective PDF export error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF report: {str(e)}")

# Export report data as PDF
@api_router.post("/export/pdf")
async def export_report_pdf(report_data: dict = Body(...), bundle: Optional[bool] = False):
    """Export report data as a downloadable PDF report"""
    logging.info("Starting PDF export for report data")
    try:
        # Helper function to clean text for PDF
        def clean_text_for_pdf(text):
            if not text:
                return ""
            text = str(text)
            # Remove emojis and special characters
            emoji_pattern = re.compile(
                "["
                "\U0001F600-\U0001F64F"  # emoticons
                "\U0001F300-\U0001F5FF"  # symbols & pictographs
                "\U0001F680-\U0001F6FF"  # transport & map symbols
                "\U0001F1E0-\U0001F1FF"  # flags (iOS)
                "\U00002702-\U000027B0"
                "\U000024C2-\U0001F251"
                "\u2022"  # bullet point
                "\u2013"  # en dash
                "\u2014"  # em dash
                "\u2018"  # left single quotation mark
                "\u2019"  # right single quotation mark
                "\u201C"  # left double quotation mark
                "\u201D"  # right double quotation mark
                "]+", flags=re.UNICODE
            )
            text = emoji_pattern.sub(r'', text)
            # Replace problematic characters
            text = text.replace('\x00', '')
            # Replace bullet points with asterisks
            text = text.replace('•', '*')
            # Replace en/em dashes with hyphens
            text = text.replace('–', '-')
            text = text.replace('—', '-')
            # Replace smart quotes with regular quotes
            text = text.replace('"', '"')
            text = text.replace('"', '"')
            text = text.replace(''', "'")
            text = text.replace(''', "'")
            return text

        logging.info("Starting PDF generation with FPDF")
        from fpdf import FPDF

        pdf = FPDF()
        # Register and prefer a Unicode-capable font when available
        try:
            dejavu_path = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
            if os.path.exists(dejavu_path):
                pdf.add_font('DejaVu', '', dejavu_path, uni=True)
                try:
                    pdf.add_font('DejaVu', 'B', dejavu_path, uni=True)
                except Exception:
                    pass
                base_font = 'DejaVu'
            else:
                base_font = 'Arial'
        except Exception:
            base_font = 'Arial'
        logging.info(f"PDF export_report_pdf: selected base_font={base_font}")

        pdf.add_page()
        pdf.set_font(base_font, 'B', 16)
        pdf.cell(0, 10, txt="CovenantAI Analysis Report", ln=True, align='C')
        pdf.ln(5)

        pdf.set_font(base_font, size=11)
        pdf.cell(0, 8, txt=f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.ln(5)

        # Document Summary
        summary_text = report_data.get('summary')
        # If structured analysis is missing but we have raw AI text, use that as summary
        if not summary_text and report_data.get('analysis') and isinstance(report_data['analysis'], dict):
            # prefer analysis.raw_text when available
            summary_text = report_data['analysis'].get('raw_text') or report_data.get('summary')

        if summary_text:
            pdf.set_font(base_font, 'B', 12)
            pdf.cell(0, 8, txt="Document Summary", ln=True)
            pdf.set_font(base_font, size=10)
            clean_content = clean_text_for_pdf(summary_text)
            logging.info(f"export_report_pdf summary length={len(summary_text)} cleaned={len(clean_content)}")
            pdf.multi_cell(0, 4, txt=clean_content)
            pdf.ln(3)

        # Risk Score
        if report_data.get('riskScore'):
            pdf.set_font(base_font, 'B', 12)
            pdf.cell(0, 8, txt="Risk Assessment Score", ln=True)
            pdf.set_font(base_font, size=10)
            risk = report_data['riskScore']
            pdf.cell(0, 6, txt=f"Score: {risk.get('score', 'N/A')}/{risk.get('max', 'N/A')}", ln=True)
            pdf.cell(0, 6, txt=f"Level: {risk.get('label', 'N/A')}", ln=True)
            pdf.ln(3)

        # SWOT Analysis
        if report_data.get('analysis'):
            analysis = report_data['analysis']

            # If analysis is not structured but contains raw_text, print the raw analysis
            if not any(k in analysis for k in ('strengths', 'weaknesses', 'opportunities', 'threats')) and analysis.get('raw_text'):
                pdf.set_font(base_font, 'B', 12)
                pdf.cell(0, 8, txt="AI Analysis", ln=True)
                pdf.set_font(base_font, size=10)
                clean_raw = clean_text_for_pdf(analysis.get('raw_text', ''))
                pdf.multi_cell(0, 4, txt=clean_raw)
                pdf.ln(2)
            else:
                # Strengths
                if analysis.get('strengths'):
                    pdf.set_font(base_font, 'B', 12)
                    pdf.cell(0, 8, txt="Strengths", ln=True)
                    pdf.set_font(base_font, size=10)
                    for item in analysis['strengths']:
                        if isinstance(item, dict) and item.get('text'):
                            clean_text = clean_text_for_pdf(item['text'])
                            pdf.multi_cell(0, 4, txt=f"* {clean_text}")
                    pdf.ln(2)

                # Weaknesses
                if analysis.get('weaknesses'):
                    pdf.set_font(base_font, 'B', 12)
                    pdf.cell(0, 8, txt="Weaknesses", ln=True)
                    pdf.set_font(base_font, size=10)
                    for item in analysis['weaknesses']:
                        if isinstance(item, dict) and item.get('text'):
                            clean_text = clean_text_for_pdf(item['text'])
                            pdf.multi_cell(0, 4, txt=f"* {clean_text}")
                    pdf.ln(2)

                # Opportunities
                if analysis.get('opportunities'):
                    pdf.set_font(base_font, 'B', 12)
                    pdf.cell(0, 8, txt="Opportunities", ln=True)
                    pdf.set_font(base_font, size=10)
                    for item in analysis['opportunities']:
                        if isinstance(item, dict) and item.get('text'):
                            clean_text = clean_text_for_pdf(item['text'])
                            pdf.multi_cell(0, 4, txt=f"* {clean_text}")
                    pdf.ln(2)

                # Threats
                if analysis.get('threats'):
                    pdf.set_font(base_font, 'B', 12)
                    pdf.cell(0, 8, txt="Threats", ln=True)
                    pdf.set_font(base_font, size=10)
                    for item in analysis['threats']:
                        if isinstance(item, dict) and item.get('text'):
                            clean_text = clean_text_for_pdf(item['text'])
                            pdf.multi_cell(0, 4, txt=f"* {clean_text}")
                    pdf.ln(2)

            # Strengths
            if analysis.get('strengths'):
                pdf.set_font(base_font, 'B', 12)
                pdf.cell(0, 8, txt="Strengths", ln=True)
                pdf.set_font(base_font, size=10)
                for item in analysis['strengths']:
                    if isinstance(item, dict) and item.get('text'):
                        clean_text = clean_text_for_pdf(item['text'])
                        pdf.multi_cell(0, 4, txt=f"* {clean_text}")
                pdf.ln(2)

            # Weaknesses
            if analysis.get('weaknesses'):
                pdf.set_font(base_font, 'B', 12)
                pdf.cell(0, 8, txt="Weaknesses", ln=True)
                pdf.set_font(base_font, size=10)
                for item in analysis['weaknesses']:
                    if isinstance(item, dict) and item.get('text'):
                        clean_text = clean_text_for_pdf(item['text'])
                        pdf.multi_cell(0, 4, txt=f"* {clean_text}")
                pdf.ln(2)

            # Opportunities
            if analysis.get('opportunities'):
                pdf.set_font(base_font, 'B', 12)
                pdf.cell(0, 8, txt="Opportunities", ln=True)
                pdf.set_font(base_font, size=10)
                for item in analysis['opportunities']:
                    if isinstance(item, dict) and item.get('text'):
                        clean_text = clean_text_for_pdf(item['text'])
                        pdf.multi_cell(0, 4, txt=f"* {clean_text}")
                pdf.ln(2)

            # Threats
            if analysis.get('threats'):
                pdf.set_font(base_font, 'B', 12)
                pdf.cell(0, 8, txt="Threats", ln=True)
                pdf.set_font(base_font, size=10)
                for item in analysis['threats']:
                    if isinstance(item, dict) and item.get('text'):
                        clean_text = clean_text_for_pdf(item['text'])
                        pdf.multi_cell(0, 4, txt=f"* {clean_text}")
                pdf.ln(2)

        # Critical Flags
        if report_data.get('criticalFlags') and len(report_data['criticalFlags']) > 0:
            pdf.set_font(base_font, 'B', 12)
            pdf.cell(0, 8, txt="Critical Red Flags", ln=True)
            pdf.set_font(base_font, size=10)
            for flag in report_data['criticalFlags']:
                if isinstance(flag, dict):
                    pdf.set_font(base_font, 'B', 10)
                    clean_title = clean_text_for_pdf(flag.get('title', ''))
                    pdf.cell(0, 6, txt=clean_title, ln=True)
                    pdf.set_font(base_font, size=9)
                    clean_explanation = clean_text_for_pdf(flag.get('explanation', ''))
                    # Include full explanation content
                    pdf.multi_cell(0, 4, txt=clean_explanation)
                    if flag.get('source'):
                        clean_source = clean_text_for_pdf(flag['source'])
                        pdf.cell(0, 4, txt=f"Source: {clean_source}", ln=True)
                    pdf.ln(2)

        # Negotiation Points
        if report_data.get('negotiationPoints') and len(report_data['negotiationPoints']) > 0:
            pdf.set_font(base_font, 'B', 12)
            pdf.cell(0, 8, txt="Negotiation Action Plan", ln=True)
            pdf.set_font(base_font, size=10)
            for point in report_data['negotiationPoints']:
                if isinstance(point, dict):
                    pdf.set_font(base_font, 'B', 10)
                    clean_title = clean_text_for_pdf(point.get('title', ''))
                    pdf.cell(0, 6, txt=clean_title, ln=True)
                    pdf.set_font(base_font, size=9)
                    if point.get('risk'):
                        clean_risk = clean_text_for_pdf(point['risk'])
                        pdf.cell(0, 4, txt=f"Risk: {clean_risk}", ln=True)
                    if point.get('example'):
                        clean_example = clean_text_for_pdf(point['example'])
                        pdf.cell(0, 4, txt=f"Suggestion: {clean_example}", ln=True)
                    pdf.ln(2)

        # Save PDF to temporary file
        logging.info("Saving PDF to temporary file")
        tmp_file_path = tempfile.mktemp(suffix='.pdf')
        pdf.output(tmp_file_path)
        logging.info(f"PDF saved successfully to: {tmp_file_path}")

        # If bundle requested, also write a UTF-8 text file with the raw analysis and return a zip
        if bundle:
            try:
                raw_text = ''
                if isinstance(report_data.get('analysis'), dict):
                    raw_text = report_data['analysis'].get('raw_text', '')
                if not raw_text:
                    # Fallback to serializing the whole report_data
                    raw_text = json.dumps(report_data, ensure_ascii=False, indent=2)

                txt_path = tempfile.mktemp(suffix='.txt')
                with open(txt_path, 'w', encoding='utf-8') as tf:
                    tf.write(raw_text)

                zip_path = tempfile.mktemp(suffix='.zip')
                import zipfile
                with zipfile.ZipFile(zip_path, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
                    zf.write(tmp_file_path, arcname='legal-analysis-report.pdf')
                    zf.write(txt_path, arcname='legal-analysis-raw.txt')

                logging.info(f"Returning bundled zip: {zip_path}")
                return FileResponse(
                    zip_path,
                    media_type='application/zip',
                    filename='legal-analysis-report.zip',
                    headers={"Content-Disposition": "attachment; filename=legal-analysis-report.zip"}
                )
            except Exception as e:
                logging.warning(f"Failed to create bundle zip: {e}")

        # Default: return PDF file
        filename = "legal-analysis-report.pdf"
        logging.info(f"Returning PDF file: {filename}")
        return FileResponse(
            tmp_file_path,
            media_type='application/pdf',
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except Exception as e:
        logging.error(f"PDF export error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF report: {str(e)}")

# Export document analysis as PDF
@api_router.get("/documents/{document_id}/export-pdf")
async def export_document_pdf(document_id: str):
    """Export document analysis as a downloadable PDF report"""
    logging.info(f"Starting PDF export for document: {document_id}")
    try:
        # Get document from database
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        document = sync_db.documents.find_one({"id": document_id})
        if not document:
            sync_client.close()
            logging.error(f"Document not found: {document_id}")
            raise HTTPException(status_code=404, detail="Document not found")

        if document.get('analysis_status') != 'completed':
            sync_client.close()
            logging.error(f"Document analysis not completed: {document_id}, status: {document.get('analysis_status')}")
            raise HTTPException(status_code=400, detail="Document analysis not completed yet")

        # Retrieve analysis sections from the document
        executive_summary = document.get('executive_summary') or ''
        plain_english = document.get('plain_english') or ''
        risk_assessment = document.get('risk_assessment') or ''
        recommendations = document.get('recommendations') or ''
        key_clauses = document.get('key_clauses') or []
        analysis_text = document.get('summary') or ''

        sync_client.close()
        logging.info(f"Retrieved document data. Summary length: {len(analysis_text)}")

        # Helper function to clean text for PDF (unicode-preserving)
        def clean_text_for_pdf(text):
            if not text:
                return ""
            text = str(text)
            # Remove NULs which can break some PDF writers
            text = text.replace('\x00', '')
            # Replace bullet points with asterisks
            text = text.replace('•', '*').replace('\u2022', '*')
            # Replace en/em dashes with hyphens
            text = text.replace('–', '-').replace('—', '-')
            # Normalize smart quotes
            text = text.replace('\u2018', "'").replace('\u2019', "'")
            text = text.replace('\u201C', '"').replace('\u201D', '"')
            return text

        logging.info("Starting PDF generation with FPDF")
        from fpdf import FPDF

        pdf = FPDF()
        # Register and prefer DejaVu for unicode support when available
        try:
            dejavu_path = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
            if os.path.exists(dejavu_path):
                pdf.add_font('DejaVu', '', dejavu_path, uni=True)
                try:
                    pdf.add_font('DejaVu', 'B', dejavu_path, uni=True)
                except Exception:
                    pass
                base_font = 'DejaVu'
            else:
                base_font = 'Arial'
        except Exception:
            base_font = 'Arial'

        pdf.add_page()
        pdf.set_font(base_font, 'B', 16)
        pdf.cell(0, 10, txt="Legal Document Analysis Report", ln=True, align='C')
        pdf.ln(5)

        pdf.set_font(base_font, size=11)
        pdf.cell(0, 8, txt=f"Document: {document.get('filename', 'Unknown')}", ln=True)
        pdf.cell(0, 8, txt=f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True)
        pdf.ln(5)

        # Add content sections
        sections = [
            ("Executive Summary", executive_summary),
            ("Plain English Summary", plain_english),
            ("Risk Assessment", risk_assessment),
            ("Recommendations", recommendations)
        ]

        for section_title, section_content in sections:
            if section_content and len(section_content.strip()) > 0:
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 8, txt=section_title, ln=True)
                pdf.set_font("Arial", size=10)
                
                clean_content = clean_text_for_pdf(section_content)
                # Allow the full section content to be written; multi_cell handles paging
                pdf.multi_cell(0, 4, txt=clean_content)
                pdf.ln(3)

    # Add key clauses if available
        if key_clauses and len(key_clauses) > 0:
            pdf.set_font("Arial", 'B', 12)
            pdf.cell(0, 8, txt="Key Clauses Analysis", ln=True)
            pdf.set_font("Arial", size=10)
            
            for clause in key_clauses:
                if clause.get('clause'):
                    clean_clause = clean_text_for_pdf(clause['clause'])
                    pdf.set_font("Arial", 'B', 10)
                    pdf.cell(0, 6, txt=f"Clause: {clean_clause[:100]}", ln=True)
                    
                    if clause.get('explanation'):
                        pdf.set_font("Arial", size=9)
                        clean_explanation = clean_text_for_pdf(clause['explanation'])
                        # Include full clause explanation content
                        pdf.multi_cell(0, 4, txt=clean_explanation)
                    pdf.ln(2)

        # Save PDF to temporary file
        logging.info("Saving PDF to temporary file")
        tmp_file_path = tempfile.mktemp(suffix='.pdf')
        pdf.output(tmp_file_path)
        
        logging.info(f"PDF saved successfully to: {tmp_file_path}")

        # Return PDF file
        filename = f"legal_analysis_{document_id}.pdf"
        logging.info(f"Returning PDF file: {filename}")
        
        return FileResponse(
            tmp_file_path,
            media_type='application/pdf',
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"PDF export error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF report: {str(e)}")

# Google OAuth endpoints
@api_router.get("/auth/google")
async def google_auth():
    """Initiate Google OAuth flow"""
    if not GOOGLE_CLIENT_ID:
        raise HTTPException(status_code=500, detail="Google OAuth not configured")

    auth_url = (
        "https://accounts.google.com/o/oauth2/v2/auth?"
        f"client_id={GOOGLE_CLIENT_ID}&"
        "response_type=code&"
        f"redirect_uri={GOOGLE_REDIRECT_URI}&"
        "scope=email%20profile&"
        "state=google_signin"
    )
    return {"auth_url": auth_url}

@api_router.get("/auth/google/callback")
async def google_auth_callback(code: str, state: str):
    """Handle Google OAuth callback"""
    if not GOOGLE_CLIENT_ID or not GOOGLE_CLIENT_SECRET:
        raise HTTPException(status_code=500, detail="Google OAuth not configured")

    # Exchange code for access token
    token_url = "https://oauth2.googleapis.com/token"
    token_data = {
        "client_id": GOOGLE_CLIENT_ID,
        "client_secret": GOOGLE_CLIENT_SECRET,
        "code": code,
        "grant_type": "authorization_code",
        "redirect_uri": GOOGLE_REDIRECT_URI,
    }

    try:
        token_response = requests.post(token_url, data=token_data)
        token_response.raise_for_status()
        token_json = token_response.json()

        # Get user info
        user_info_url = "https://www.googleapis.com/oauth2/v2/userinfo"
        headers = {"Authorization": f"Bearer {token_json['access_token']}"}
        user_response = requests.get(user_info_url, headers=headers)
        user_response.raise_for_status()
        user_info = user_response.json()

        # Check if user exists, create if not
        sync_client = pymongo.MongoClient(os.environ.get('MONGO_URL', 'mongodb://localhost:27017'))
        sync_db = sync_client[os.environ.get('DB_NAME', 'legal_docs')]
        users_collection = sync_db.users

        existing_user = users_collection.find_one({"google_id": user_info["id"]})
        if not existing_user:
            user_doc = {
                "id": str(uuid.uuid4()),
                "email": user_info["email"],
                "name": user_info["name"],
                "google_id": user_info["id"],
                "created_at": datetime.now(timezone.utc)
            }
            users_collection.insert_one(user_doc)
            user = user_doc
        else:
            user = existing_user

        sync_client.close()

        # Create JWT token
        access_token = create_access_token(
            data={"sub": user["id"], "email": user["email"]},
            expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        )

        return {
            "access_token": access_token,
            "token_type": "bearer",
            "user": {
                "id": user["id"],
                "email": user["email"],
                "name": user["name"]
            }
        }

    except requests.RequestException as e:
        logging.error(f"OAuth error: {e}")
        raise HTTPException(status_code=400, detail="OAuth authentication failed")

# Then include the router in the main app
app.include_router(api_router)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    if client:
        client.close()

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
